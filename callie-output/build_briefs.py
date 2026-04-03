#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = "/Users/minigrill/.openclaw/workspace/callie-output"

FIELDS = [
    "Email Objective",
    "Subject Line Directions",
    "Preview Text Direction",
    "Primary Message",
    "Content Block Structure",
    "CTA Direction",
    "Tone Notes",
]

# ── colour palette ──────────────────────────────────────────────
BLACK      = RGBColor(0x1A, 0x1A, 0x1A)
DARK_GREY  = RGBColor(0x33, 0x33, 0x33)
MID_GREY   = RGBColor(0x66, 0x66, 0x66)
LABEL_BG   = RGBColor(0xF2, 0xF2, 0xF2)  # light grey fill for label column
HEADING_BG = RGBColor(0x1A, 0x1A, 0x1A)  # near-black for email headings
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, top=True, bottom=True, left=True, right=True,
                     color="DDDDDD", sz="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, active in [("top", top), ("bottom", bottom),
                          ("left", left), ("right", right)]:
        el = OxmlElement(f"w:{side}")
        if active:
            el.set(qn("w:val"),   "single")
            el.set(qn("w:sz"),    sz)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
        else:
            el.set(qn("w:val"), "none")
        tcBorders.append(el)
    tcPr.append(tcBorders)

def add_paragraph_spacing(doc, space_before=0, space_after=6):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    return p

def add_horizontal_rule(doc):
    """Thin grey rule between email sections."""
    p  = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(10)
    pf.space_after  = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def build_doc(persona_name, persona_tag, emails):
    """
    persona_name : e.g. "The High-Performance Executive"
    persona_tag  : e.g. "P1 High-Performance Executive"
    emails       : list of dicts with keys matching FIELDS + 'heading'
    """
    doc = Document()

    # ── page margins ────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── default body font ────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)
    style.font.color.rgb = DARK_GREY

    # ── TITLE ────────────────────────────────────────────────────
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(0)
    title_p.paragraph_format.space_after  = Pt(4)
    run = title_p.add_run("Rho Nutrition")
    run.bold       = True
    run.font.size  = Pt(26)
    run.font.color.rgb = BLACK
    run.font.name  = "Calibri"

    # ── SUBTITLE LINE 1 ──────────────────────────────────────────
    sub1 = doc.add_paragraph()
    sub1.paragraph_format.space_before = Pt(0)
    sub1.paragraph_format.space_after  = Pt(2)
    r = sub1.add_run(f"Post Purchase Flow Briefs  |  {persona_name}")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = DARK_GREY
    r.font.name = "Calibri"

    # ── SUBTITLE LINE 2 ──────────────────────────────────────────
    sub2 = doc.add_paragraph()
    sub2.paragraph_format.space_before = Pt(0)
    sub2.paragraph_format.space_after  = Pt(2)
    r2 = sub2.add_run("Emails 1, 2 + 3  |  Day 3 / Day 8 / Day 21")
    r2.font.size = Pt(11)
    r2.font.color.rgb = MID_GREY
    r2.font.name = "Calibri"

    # ── SUBTITLE LINE 3 ──────────────────────────────────────────
    sub3 = doc.add_paragraph()
    sub3.paragraph_format.space_before = Pt(0)
    sub3.paragraph_format.space_after  = Pt(20)
    r3 = sub3.add_run("DigitsUp x Rho Nutrition  |  Outersignal Project")
    r3.font.size = Pt(10)
    r3.font.color.rgb = MID_GREY
    r3.italic = True
    r3.font.name = "Calibri"

    # ── EMAILS ───────────────────────────────────────────────────
    for idx, email in enumerate(emails):
        # Separator (not before the first email)
        if idx > 0:
            add_horizontal_rule(doc)
            add_paragraph_spacing(doc, space_before=4, space_after=4)

        # Email heading paragraph (dark background)
        heading_p = doc.add_paragraph()
        heading_p.paragraph_format.space_before = Pt(6)
        heading_p.paragraph_format.space_after  = Pt(12)

        # Use a shaded table row trick for the heading banner
        heading_tbl = doc.add_table(rows=1, cols=1)
        heading_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        heading_tbl.style = "Table Grid"
        h_cell = heading_tbl.cell(0, 0)
        set_cell_bg(h_cell, HEADING_BG)
        # remove inner borders
        set_cell_borders(h_cell, color="1A1A1A")
        h_cell._tc.get_or_add_tcPr()
        h_para = h_cell.paragraphs[0]
        h_para.paragraph_format.space_before = Pt(6)
        h_para.paragraph_format.space_after  = Pt(6)
        h_para.paragraph_format.left_indent  = Pt(8)
        h_run = h_para.add_run(email["heading"])
        h_run.bold = True
        h_run.font.size = Pt(11)
        h_run.font.color.rgb = WHITE
        h_run.font.name = "Calibri"

        # Remove the spurious empty paragraph added before the table
        heading_p._element.getparent().remove(heading_p._element)

        # spacer
        sp = doc.add_paragraph()
        sp.paragraph_format.space_before = Pt(0)
        sp.paragraph_format.space_after  = Pt(6)

        # Content table
        tbl = doc.add_table(rows=len(FIELDS), cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        tbl.style = "Table Grid"

        # Column widths  (label ≈ 5 cm, content ≈ 12 cm)
        for row_idx, field in enumerate(FIELDS):
            row = tbl.rows[row_idx]
            row.height_rule = None

            label_cell   = row.cells[0]
            content_cell = row.cells[1]

            # Set column widths via XML
            label_cell._tc.get_or_add_tcPr()
            content_cell._tc.get_or_add_tcPr()

            # Background on label cell
            set_cell_bg(label_cell, LABEL_BG)

            # Borders
            border_color = "CCCCCC"
            set_cell_borders(label_cell,   color=border_color)
            set_cell_borders(content_cell, color=border_color)

            # Label text
            lp = label_cell.paragraphs[0]
            lp.paragraph_format.space_before = Pt(4)
            lp.paragraph_format.space_after  = Pt(4)
            lp.paragraph_format.left_indent  = Pt(6)
            lr = lp.add_run(field)
            lr.bold = True
            lr.font.size = Pt(9.5)
            lr.font.color.rgb = BLACK
            lr.font.name = "Calibri"

            # Content text
            cp = content_cell.paragraphs[0]
            cp.paragraph_format.space_before = Pt(4)
            cp.paragraph_format.space_after  = Pt(4)
            cp.paragraph_format.left_indent  = Pt(6)
            cp.paragraph_format.right_indent = Pt(6)
            cr = cp.add_run(email.get(field, ""))
            cr.font.size = Pt(9.5)
            cr.font.color.rgb = DARK_GREY
            cr.font.name = "Calibri"

        # Set column widths explicitly
        for row in tbl.rows:
            row.cells[0].width = Inches(1.9)
            row.cells[1].width = Inches(4.7)

        # Trailing space
        sp2 = doc.add_paragraph()
        sp2.paragraph_format.space_before = Pt(0)
        sp2.paragraph_format.space_after  = Pt(8)

    return doc


# ═══════════════════════════════════════════════════════════════
#  BRIEF CONTENT
# ═══════════════════════════════════════════════════════════════

BRIEFS = [

    # ── P1 ──────────────────────────────────────────────────────
    {
        "persona_name": "The High-Performance Executive",
        "persona_tag":  "P1 High-Performance Executive",
        "emails": [
            {
                "heading": "EMAIL 1 — DAY 3  |  THE HIGH-PERFORMANCE EXECUTIVE",
                "Email Objective": "Reinforce the purchase decision and frame daily usage as a performance protocol. Remove friction from starting. Position Rho as a high-ROI addition to an already optimized life.",
                "Subject Line Directions": "Your optimization protocol starts now / One serving. Every morning. Here is why it matters. / You made a smart call. Here is how to get the most from it.",
                "Preview Text Direction": "Short, efficient. Mirror their mindset. Example: 'Takes 30 seconds. Works all day.'",
                "Primary Message": "You invested in your performance. Now activate it. One serving of Rho in the morning delivers NAD+ and liposomal nutrients at the cellular level. Faster than capsules. No wasted product. Frame as a protocol, not a routine.",
                "Content Block Structure": "Hero: Confident welcome. 'Your optimization protocol starts now.' Acknowledge the purchase as a smart decision.\n\nSecondary: Usage instructions — broad, applicable across all Rho products. One serving, morning timing, consistency. Keep it clean and protocol-style, not product-specific.\n\nSecondary: Brief liposomal differentiator (why it absorbs better). No more than 2-3 sentences.",
                "CTA Direction": "Direct and action-oriented. 'Add to Your Morning Stack.' No urgency. No discount. Just the next logical step.",
                "Tone Notes": "Boardroom confidence. Zero fluff. Short sentences. Every line earns its place. Think internal memo, not wellness newsletter.",
            },
            {
                "heading": "EMAIL 2 — DAY 8  |  THE HIGH-PERFORMANCE EXECUTIVE",
                "Email Objective": "Validate the decision with data. One week in, explain what NAD+ is doing at the cellular level. Build belief through science, not testimonials.",
                "Subject Line Directions": "What NAD+ is doing right now / The cellular results you cannot see yet / One week in. Here is what is happening.",
                "Preview Text Direction": "Direct, data-led. Example: \"Cellular energy, cognitive clarity, compounding results. The science behind what you are feeling.\"",
                "Primary Message": "At the cellular level, NAD+ is repairing DNA, fueling mitochondria, and supporting cognitive function. The results compound. Explain the mechanism briefly, confirm they are ahead of most people, and reinforce daily consistency as the performance variable.",
                "Content Block Structure": "Hero: Week one validation, data-led. Brief acknowledgment of decision quality.\n\nSecondary: What NAD+ does at the cellular level. 2-3 sentences, precise. Mitochondrial function, DNA repair, cognitive support.\n\nSecondary: Liposomal delivery vs capsules. Why absorption matters and what they chose over standard supplements.\n\nSecondary: Consistency framing. Results compound. Do not stop.",
                "CTA Direction": "\"Read the Science\" or \"See the Research\" — evidence-driven, no urgency.",
                "Tone Notes": "Data-led, efficient, precise. No enthusiasm. No wellness language. Boardroom voice throughout.",
            },
            {
                "heading": "EMAIL 3 — DAY 21  |  THE HIGH-PERFORMANCE EXECUTIVE",
                "Email Objective": "Frame consistency as a protocol, not a habit. Cost-per-day logic. Prevent a gap in the stack.",
                "Subject Line Directions": "Do not let your results reset / The protocol only works if you do not stop / Your next supply. Before you need it.",
                "Preview Text Direction": "Logical and outcome-focused. Example: \"Three weeks of cellular optimization. The ROI compounds only if the protocol continues.\"",
                "Primary Message": "Three weeks of consistent NAD+ and liposomal nutrition is the baseline, not the destination. The results compound with time. Stopping resets the cellular work. Cost-per-day framing: compare to a single business lunch. The protocol is the competitive edge. Do not interrupt it.",
                "Content Block Structure": "Hero: Three weeks in. Brief data-driven acknowledgment of progress. Tone: boardroom, not celebratory.\n\nDYNAMIC BLOCK (Secondary):\n\nSUBSCRIBER VERSION: Acknowledge they are locked in. Reinforce the decision with one line about compounding results. CTA to account dashboard.\n\nNON-SUBSCRIBER VERSION: Cost-per-day framing. One serving per day. Compare to the cost of a business lunch, a coffee meeting, a single supplement that does not absorb. Subscription = logical protocol decision. Skip or cancel anytime.",
                "CTA Direction": "\"Secure Your Next Supply\" (non-subscriber) / \"You Are Locked In\" (subscriber)",
                "Tone Notes": "Logical, outcome-focused. Boardroom throughout. Zero sentiment.",
            },
        ],
    },

    # ── P2 ──────────────────────────────────────────────────────
    {
        "persona_name": "The Creative Entrepreneur",
        "persona_tag":  "P2 Creative Entrepreneur",
        "emails": [
            {
                "heading": "EMAIL 1 — DAY 3  |  THE CREATIVE ENTREPRENEUR",
                "Email Objective": "Welcome the customer into the Rho ritual. Frame daily usage as an intentional, elevated moment. Make starting feel inspiring, not clinical.",
                "Subject Line Directions": "Your new ritual is ready / The morning routine just got an upgrade / Clean. Simple. Yours.",
                "Preview Text Direction": "Warm and sensory. Example: 'A 30-second ritual that sets the tone for everything after.'",
                "Primary Message": "You chose a product that fits the way you live. Clean ingredients, thoughtful formulation, nothing unnecessary. The liquid format is intentional. Position this as a small act of self-care that supports creative output. Frame as a daily ritual, not a supplement routine.",
                "Content Block Structure": "Hero: Aesthetic welcome. Clean, morning energy. Warm tone. Not clinical.\n\nSecondary: Usage instructions — broad, universal. One serving in the morning. Frame it as a ritual moment, not a clinical protocol. Keep product-specific details out; the tone does the heavy lifting.\n\nSecondary: Clean label callout. What is inside. What is not. Brief and confident.",
                "CTA Direction": "Soft and inviting. 'Make It Your Morning Ritual.' Feels like an invitation, not an instruction.",
                "Tone Notes": "Warm, considered, artisan. Reads like a brand they already love. No clinical language. No mass-market wellness. White space is your friend.",
            },
            {
                "heading": "EMAIL 2 — DAY 8  |  THE CREATIVE ENTREPRENEUR",
                "Email Objective": "Connect clean ingredients to mental clarity and creative output. One week in — validate what they may be noticing. Build ingredient transparency as a brand value.",
                "Subject Line Directions": "Clean ingredients for a clear mind / One week in. What are you noticing? / The label that says exactly what it is.",
                "Preview Text Direction": "Warm and reflective. Example: \"What is inside, why it matters, and what a week of clean nutrition can quietly shift.\"",
                "Primary Message": "Clean label. No fillers, no synthetics, no surprises. Walk through the key ingredients in a way that feels like brand storytelling, not a clinical breakdown. NAD+ for mental sharpness. Glutathione for cellular clarity. It fits the way they live.",
                "Content Block Structure": "Hero: One week check-in. Warm, curious. \"What are you noticing?\" Gentle, not clinical.\n\nSecondary: Ingredient transparency. Key compounds named, one line each, lifestyle framing. NAD+, Glutathione, Vitamin C.\n\nSecondary: Clean label philosophy. What is not in Rho is as important as what is. No fillers, no synthetics.\n\nSecondary: Gentle consistency note. The ritual is working. Keep it.",
                "CTA Direction": "\"See What Is Inside\" — feels like an invitation, not an instruction.",
                "Tone Notes": "Considered, honest, artisan. Reads like a brand they already love. No clinical language. No mass-market wellness.",
            },
            {
                "heading": "EMAIL 3 — DAY 21  |  THE CREATIVE ENTREPRENEUR",
                "Email Objective": "Protect the ritual. Frame subscription as never having to think about it again. Identity-affirming.",
                "Subject Line Directions": "Keep the ritual going / Your routine is worth protecting / Never run out of what is working",
                "Preview Text Direction": "Gentle and identity-affirming. Example: \"Three weeks of something that is just yours. Here is how to keep it.\"",
                "Primary Message": "The ritual is working. Do not let it lapse because of logistics. Subscription is the simplest version of self-care: the thing that works shows up when it should. Frame it as removing friction from something worth protecting.",
                "Content Block Structure": "Hero: Three weeks in. Warm acknowledgment. \"This has become part of how you start your day.\"\n\nDYNAMIC BLOCK (Secondary):\n\nSUBSCRIBER VERSION: Warm affirmation. You are already set. The ritual continues effortlessly. Reinforce the decision gently. CTA to account.\n\nNON-SUBSCRIBER VERSION: Subscription as effortless continuity. Never run out. Never have to think about it. Your ritual, on autopilot. Skip or cancel anytime. Feels like looking after yourself, not a transaction.",
                "CTA Direction": "\"Never Run Out\" (non-subscriber) / \"Your Ritual Continues\" (subscriber)",
                "Tone Notes": "Warm, gentle, identity-affirming. No urgency. No discount language.",
            },
        ],
    },

    # ── P3 ──────────────────────────────────────────────────────
    {
        "persona_name": "The Community Educator and Parent",
        "persona_tag":  "P3 Community Educator and Parent",
        "emails": [
            {
                "heading": "EMAIL 1 — DAY 3  |  THE COMMUNITY EDUCATOR AND PARENT",
                "Email Objective": "Make the customer feel good about their decision. Reassuring usage instructions. Reinforce that this is a safe, trusted product built for real people with full lives.",
                "Subject Line Directions": "You made a good choice for yourself / Here is how to get started with Rho / Simple steps for your new daily routine",
                "Preview Text Direction": "Warm and affirming. Example: 'You deserve this. Here is how to make it work for your day.'",
                "Primary Message": "Between students, family, and everything else they carry, this is one thing that is just for them. Lead with warmth. Reinforce safety, simplicity, and the founder mission. One serving. That is all.",
                "Content Block Structure": "Hero: Warm welcome. Acknowledge how much they give. This is for them.\n\nSecondary: Simple usage instructions — broad and universal. One serving in the morning, with or without food. No jargon. Tone is warm and encouraging, not clinical or product-specific.\n\nSecondary: Safety and trust block. Clean ingredients. Founder mission (2 sentences). Built for real families.",
                "CTA Direction": "Encouraging and low-pressure. 'Start Your Routine Today.' Feels like a nudge from a trusted friend.",
                "Tone Notes": "Human, warm, reassuring. No science-heavy language. No performance framing. This persona responds to trust and community.",
            },
            {
                "heading": "EMAIL 2 — DAY 8  |  THE COMMUNITY EDUCATOR AND PARENT",
                "Email Objective": "Reinforce immune support and safety. One week in — acknowledge the high-demand environment they work in. Build trust through ingredient transparency and founder mission.",
                "Subject Line Directions": "Immune support built for your world / One week in. Keep going. / The ingredients you can trust.",
                "Preview Text Direction": "Warm and affirming. Example: \"You are around a lot of people every day. Here is what Rho is doing on your behalf.\"",
                "Primary Message": "High-contact environments demand stronger immune support. Glutathione is the body's master antioxidant. Vitamin C supports immune response. This is not about performance. It is about staying well so you can keep showing up. Founder mission reinforces trust.",
                "Content Block Structure": "Hero: One week warm check-in. \"You are doing it. Keep going.\" Acknowledges effort, not performance.\n\nSecondary: Immune support callout. Glutathione + Vitamin C, explained simply and warmly. No jargon.\n\nSecondary: Clean and safe ingredients. No harmful additives. Founder story in 2 sentences. Trust-building.\n\nSecondary: Gentle consistency encouragement.",
                "CTA Direction": "\"Learn What Is Inside\" — warm, reassuring, low-pressure.",
                "Tone Notes": "Community, care, trust. No performance framing. Simple language throughout. Human first.",
            },
            {
                "heading": "EMAIL 3 — DAY 21  |  THE COMMUNITY EDUCATOR AND PARENT",
                "Email Objective": "Reliability for the people who count on them. Frame consistency as showing up for others. Low-pressure.",
                "Subject Line Directions": "Stay consistent for the people who count on you / Three weeks in. Keep going. / Your routine is working. Do not let it lapse.",
                "Preview Text Direction": "Warm and community-centered. Example: \"You show up for everyone else. This is one way to show up for yourself, consistently.\"",
                "Primary Message": "Three weeks is real progress. Immune support and sustained energy do not stop building. They compound with time. The people around you benefit when you stay consistent. Keep the supply going. This is not a luxury. It is maintenance.",
                "Content Block Structure": "Hero: Warm three-week acknowledgment. \"You have been consistent. That matters.\"\n\nDYNAMIC BLOCK (Secondary):\n\nSUBSCRIBER VERSION: Warm affirmation. Reinforce they are already set. Thank them. CTA to account.\n\nNON-SUBSCRIBER VERSION: Never run out framing. Skip or cancel note. Framed around reliability and showing up for others. Warm, not transactional.",
                "CTA Direction": "\"Keep Your Supply Going\" (non-subscriber) / \"You Are All Set\" (subscriber)",
                "Tone Notes": "Nurturing, community-first. No performance language. No urgency. Simple and warm throughout.",
            },
        ],
    },

    # ── P4 ──────────────────────────────────────────────────────
    {
        "persona_name": "The Clinical Wellness Advocate",
        "persona_tag":  "P4 Clinical Wellness Advocate",
        "emails": [
            {
                "heading": "EMAIL 1 — DAY 3  |  THE CLINICAL WELLNESS ADVOCATE",
                "Email Objective": "Confirm the clinical rationale for their purchase. Provide precise dosing and timing guidance. Acknowledge their expertise. Confirm they made the right call.",
                "Subject Line Directions": "Your liposomal protocol is active / Dosing, timing, and what to expect from day one / The formulation you chose. Here is how to use it correctly.",
                "Preview Text Direction": "Clinical and direct. Example: 'Optimal absorption window, full formulation breakdown, and what to monitor.'",
                "Primary Message": "They understand bioavailability. They chose liposomal for a reason. Confirm they were right and give them the clinical detail they expect: dosing, timing, absorption, key compounds. Peer-level throughout.",
                "Content Block Structure": "Hero: Peer-level acknowledgment. 'You already know what liposomal means. Here is how to use Rho correctly.'\n\nSecondary: Dosing and timing — broad and universal. Morning timing, consistency, absorption window. Do not prescribe product-specific instructions since formulations vary. Stay at the protocol level.\n\nSecondary: Liposomal delivery brief callout. 2 sentences. Correct terminology.",
                "CTA Direction": "Evidence-driven. 'View the Full Formulation.' Invites them deeper without overselling.",
                "Tone Notes": "Peer-level. Use correct terminology: bioavailability, liposomal, phospholipid, cellular absorption. Confident and precise. Never dumbed down. Never enthusiastic.",
            },
            {
                "heading": "EMAIL 2 — DAY 8  |  THE CLINICAL WELLNESS ADVOCATE",
                "Email Objective": "Go deep on the bioavailability science. Peer-level education on liposomal delivery vs standard oral supplements. Validate their recommendation confidence.",
                "Subject Line Directions": "The bioavailability difference, by the numbers / Liposomal delivery: what the research shows / One week in. The mechanism is working.",
                "Preview Text Direction": "Clinical and precise. Example: \"Phospholipid encapsulation, absorption rates, and why liposomal is the standard your patients deserve.\"",
                "Primary Message": "Standard oral supplements lose significant potency before reaching target cells. Liposomal encapsulation using phospholipid bilayers protects active compounds through the GI tract and delivers them at the cellular level. The data is clear. This is the formulation worth recommending.",
                "Content Block Structure": "Hero: Peer acknowledgment. \"One week in. The mechanism is active.\" Clinical, no congratulations.\n\nSecondary: Liposomal vs standard. Absorption comparison, mechanism explained with correct terminology: phospholipid bilayer, hepatic bypass, cellular uptake.\n\nSecondary: Key compounds and clinical significance. Brief, precise, named correctly.\n\nSecondary: Consistency framing. Uninterrupted dosing maintains cellular saturation.",
                "CTA Direction": "\"Read the Research\" — evidence-driven, peer-appropriate.",
                "Tone Notes": "Peer-level throughout. Correct terminology only. No simplification. No enthusiasm. Precise and confident.",
            },
            {
                "heading": "EMAIL 3 — DAY 21  |  THE CLINICAL WELLNESS ADVOCATE",
                "Email Objective": "Clinical consistency as protocol. Uninterrupted supplementation produces measurable outcomes. Professional, outcome-driven.",
                "Subject Line Directions": "Clinical consistency drives clinical results / Do not interrupt the protocol / Three weeks in. Maintain the regimen.",
                "Preview Text Direction": "Clinical and direct. Example: \"Interrupted supplementation resets cellular baseline. Maintaining the protocol is the clinical decision.\"",
                "Primary Message": "Liposomal supplementation requires consistent dosing to maintain cellular saturation. A gap in the protocol means reestablishing baseline. For patients you may recommend this to, the message is the same: uninterrupted use produces measurable outcomes. Maintain the protocol.",
                "Content Block Structure": "Hero: Clinical three-week milestone. Brief acknowledgment. Tone: peer, not congratulatory.\n\nDYNAMIC BLOCK (Secondary):\n\nSUBSCRIBER VERSION: Acknowledge protocol maintained. Confirm cellular work is uninterrupted. CTA to account dashboard.\n\nNON-SUBSCRIBER VERSION: Interrupted supplementation data point. What resets when dosing stops. Subscription = protocol continuity. Clinical rationale for uninterrupted supply. Skip or cancel framed clinically.",
                "CTA Direction": "\"Maintain Your Protocol\" (non-subscriber) / \"Protocol Maintained\" (subscriber)",
                "Tone Notes": "Professional, clinical, outcome-driven. No warmth. No lifestyle framing. Peer-level throughout.",
            },
        ],
    },

    # ── P5 ──────────────────────────────────────────────────────
    {
        "persona_name": "The Active Performance Optimizer",
        "persona_tag":  "P5 Active Performance Optimizer",
        "emails": [
            {
                "heading": "EMAIL 1 — DAY 3  |  THE ACTIVE PERFORMANCE OPTIMIZER",
                "Email Objective": "Connect Rho directly to their training and recovery routine. Frame day one as the start of a performance system. Give them a clear, actionable stacking protocol.",
                "Subject Line Directions": "Recovery starts with day one / Here is how to stack Rho into your routine / Day one. The protocol begins.",
                "Preview Text Direction": "Action-oriented. Example: 'Morning timing, anti-inflammatory support, and what you should start noticing by week two.'",
                "Primary Message": "Recovery is not what happens after training. It is the system built before the next session. Rho fits into that system. Morning timing, consistent daily use, Curcumin and Resveratrol doing their work.",
                "Content Block Structure": "Hero: Performance framing. 'Recovery does not start after the session. It starts here.'\n\nSecondary: Stacking protocol — broad and universal. When to take it, why morning timing matters for anti-inflammatory support. Do not prescribe product-specific usage since formulations vary across the Rho line. Keep it at the system level.\n\nSecondary: Brief ingredient callout. Curcumin + Resveratrol for recovery framing. 2-3 sentences.",
                "CTA Direction": "Active and peer-level. 'Add It to Your Stack.' Sounds like advice from a trusted coach, not a brand.",
                "Tone Notes": "Knowledgeable, direct, energetic. Results and recovery are the only currencies. No lifestyle language. No ritual framing.",
            },
            {
                "heading": "EMAIL 2 — DAY 8  |  THE ACTIVE PERFORMANCE OPTIMIZER",
                "Email Objective": "Validate the physical signals they may be noticing. Explain what Curcumin and Resveratrol are doing inside joints and muscles. Build belief through mechanism, not testimonial.",
                "Subject Line Directions": "What Curcumin is doing inside your joints / One week in. Here is what is happening. / The anti-inflammatory science behind your recovery.",
                "Preview Text Direction": "Active and direct. Example: \"Curcumin, Resveratrol, inflammation pathways, and what consistent use means for your next session.\"",
                "Primary Message": "Curcumin inhibits inflammatory pathways at the molecular level. Resveratrol supports mitochondrial function and reduces oxidative stress. In a liposomal format, these compounds actually reach the tissues where you need them. One week is not the peak. It is the start of compounding recovery.",
                "Content Block Structure": "Hero: One week in validation. \"You might be noticing something. Here is the science behind it.\" Direct, no fluff.\n\nSecondary: Curcumin mechanism. NF-kB pathway, joint inflammation, why liposomal matters for bioavailability. Performance-first framing.\n\nSecondary: Resveratrol. Mitochondrial function, oxidative stress, energy for training.\n\nSecondary: Consistency framing. Week four is where compounding becomes measurable.",
                "CTA Direction": "\"See the Science Behind Recovery\" — active and peer-level.",
                "Tone Notes": "Knowledgeable, energetic, direct. Results and recovery are the only currencies. No lifestyle language.",
            },
            {
                "heading": "EMAIL 3 — DAY 21  |  THE ACTIVE PERFORMANCE OPTIMIZER",
                "Email Objective": "Consistency is compounding recovery. Do not lose the progress built over three weeks. Coach-like energy.",
                "Subject Line Directions": "Do not lose the progress you have built / Three weeks in. The compounding starts now. / Stay in the stack.",
                "Preview Text Direction": "Motivating and direct. Example: \"Three weeks of Curcumin and Resveratrol building in your system. Week four is where it compounds.\"",
                "Primary Message": "Three weeks of consistent anti-inflammatory support is not the peak. It is the foundation. The compounds are building in your system. Week four and beyond is where compounding recovery becomes measurable. A gap means starting over. Stay in the stack.",
                "Content Block Structure": "Hero: Three weeks. Performance validation. \"You are building something. Do not stop now.\"\n\nDYNAMIC BLOCK (Secondary):\n\nSUBSCRIBER VERSION: You are locked in. Compounding continues. Brief acknowledgment. CTA to account.\n\nNON-SUBSCRIBER VERSION: What stops when you stop. Curcumin and Resveratrol leave the system. The progress resets. Subscription = staying in the stack without thinking about it. Skip or cancel anytime.",
                "CTA Direction": "\"Stay in the Stack\" (non-subscriber) / \"Locked In. Keep Going.\" (subscriber)",
                "Tone Notes": "Coach-like, motivating, direct. Performance language only. No lifestyle or ritual framing.",
            },
        ],
    },

    # ── P6 ──────────────────────────────────────────────────────
    {
        "persona_name": "The Skilled Trade and Industry Leader",
        "persona_tag":  "P6 Skilled Trade and Industry Leader",
        "emails": [
            {
                "heading": "EMAIL 1 — DAY 3  |  THE SKILLED TRADE AND INDUSTRY LEADER",
                "Email Objective": "Get straight to it. How to take it and what to expect. No story. No ritual. Practical instructions from someone who respects their time.",
                "Subject Line Directions": "Here is how to get the most out of it / Simple. Consistent. Effective. Your guide to Rho. / You bought it. Here is exactly how to use it.",
                "Preview Text Direction": "Blunt and practical. Example: 'One serving a day. Morning is best. Here is what consistent use does.'",
                "Primary Message": "Take one serving in the morning. Do it every day. Explain liposomal vs standard in one paragraph using money language. Joint support and energy take weeks. No hype. Stick with it.",
                "Content Block Structure": "Hero: No-fluff welcome. 'You made a solid decision. Here is how to use it right.'\n\nSecondary: Usage instructions — broad and universal. Bullet format. One serving, morning timing, with or without food. No product-specific detail. Blunt and practical.\n\nSecondary: Why liposomal works. One short paragraph. Money language: your money is actually absorbing. No jargon.",
                "CTA Direction": "Blunt and direct. 'Start Today.' No decoration.",
                "Tone Notes": "Straight talk. Short sentences. Zero lifestyle or wellness language. If it sounds like a wellness newsletter, rewrite it.",
            },
            {
                "heading": "EMAIL 2 — DAY 8  |  THE SKILLED TRADE AND INDUSTRY LEADER",
                "Email Objective": "Explain why this works when other supplements did not. Liposomal vs standard in plain money language. Validate joint health support for physical demands. Keep it short and credible.",
                "Subject Line Directions": "Why this works when others do not / One week in. Here is what to expect next. / The difference you paid for, explained plainly.",
                "Preview Text Direction": "Blunt and practical. Example: \"Most supplements never make it to where you need them. Here is why Rho does.\"",
                "Primary Message": "Most supplements pass through without absorbing. Liposomal means the nutrients actually get into your bloodstream. Joint support from Curcumin takes a few weeks to build. That is normal. Keep going. You are not wasting your money.",
                "Content Block Structure": "Hero: Straight talk. \"One week in. Here is what is happening.\" Short. No buildup.\n\nSecondary: Liposomal explained simply. One short paragraph. Framed as: your money is actually working. No scientific jargon.\n\nSecondary: Joint health timeline. Curcumin takes 2-4 weeks to feel it. Honest, no hype. Sets expectation.\n\nSecondary: Keep going. Short. Blunt.",
                "CTA Direction": "\"See Why It Is Different\" — plain, direct.",
                "Tone Notes": "Straight, credible, practical. Zero lifestyle language. Short sentences. No wellness newsletter energy.",
            },
            {
                "heading": "EMAIL 3 — DAY 21  |  THE SKILLED TRADE AND INDUSTRY LEADER",
                "Email Objective": "Running low? Simple. Do not let it run out. No restart needed. Lock it in.",
                "Subject Line Directions": "Running low? Do not let it run out. / Three weeks in. Lock it in. / Simple. Keep it going.",
                "Preview Text Direction": "Blunt and practical. Example: \"You are three weeks in. Do not start over.\"",
                "Primary Message": "You are three weeks in. The joint support is building. Do not let it run out and start over. Lock in a subscription and forget about it. Skip or cancel anytime. Just do not let it reset.",
                "Content Block Structure": "Hero: Three weeks. Short, direct. \"Good. Keep going.\"\n\nDYNAMIC BLOCK (Secondary):\n\nSUBSCRIBER VERSION: You are sorted. Keep going. CTA to account. Two sentences maximum.\n\nNON-SUBSCRIBER VERSION: Do not start over framing. Subscription means locked in. Skip or cancel note. Cost-per-day in plain language. No fluff.",
                "CTA Direction": "\"Lock It In\" (non-subscriber) / \"You Are Sorted\" (subscriber)",
                "Tone Notes": "Blunt, short, zero decoration. Straight talk. If it sounds like a wellness newsletter, rewrite it.",
            },
        ],
    },
]

FILE_NAMES = [
    "Rho Nutrition Brief — P1 High-Performance Executive.docx",
    "Rho Nutrition Brief — P2 Creative Entrepreneur.docx",
    "Rho Nutrition Brief — P3 Community Educator and Parent.docx",
    "Rho Nutrition Brief — P4 Clinical Wellness Advocate.docx",
    "Rho Nutrition Brief — P5 Active Performance Optimizer.docx",
    "Rho Nutrition Brief — P6 Skilled Trade and Industry Leader.docx",
]

if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    for brief, fname in zip(BRIEFS, FILE_NAMES):
        doc = build_doc(brief["persona_name"], brief["persona_tag"], brief["emails"])
        path = os.path.join(OUTPUT_DIR, fname)
        doc.save(path)
        print(path)
    print("\nDone. 6 files saved.")
