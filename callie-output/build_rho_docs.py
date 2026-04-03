#!/usr/bin/env python3
"""Build 6 Rho Nutrition post-purchase email copy .docx files."""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = "/Users/minigrill/.openclaw/workspace/callie-output"

# ── helpers ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_paragraph_in_cell(cell, text, bold=False, size=10, space_before=0):
    """Add a paragraph to a cell, returning the paragraph."""
    p = cell.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p

def fill_cell_body(cell, text, size=10):
    """Fill a cell with multi-paragraph body text (split on \\n)."""
    paragraphs = text.split('\n')
    first = True
    for para_text in paragraphs:
        if first:
            # Use the existing empty paragraph in the cell
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(para_text)
        run.font.size = Pt(size)

def add_email_table(doc, email_num, day, persona_name, subject, preview, body, cta):
    """Add a 2-col table for one email."""
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'

    # Column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.5)
        row.cells[1].width = Inches(5.0)

    fields = [
        ("Email", f"EMAIL {email_num} — DAY {day}"),
        ("Persona", persona_name),
        ("Subject Line", subject),
        ("Preview Text", preview),
        ("Body Copy", None),  # handled separately
        ("CTA", cta),
    ]

    for i, (label, value) in enumerate(fields):
        lc = table.rows[i].cells[0]
        vc = table.rows[i].cells[1]
        set_cell_bg(lc, "F5F5F5")

        # Label
        lc.paragraphs[0].clear()
        run = lc.paragraphs[0].add_run(label)
        run.bold = True
        run.font.size = Pt(10)

        # Value
        if label == "Body Copy":
            fill_cell_body(vc, body)
        else:
            vc.paragraphs[0].clear()
            run2 = vc.paragraphs[0].add_run(value)
            run2.font.size = Pt(10)
            if label == "CTA":
                run2.bold = True

    doc.add_paragraph()  # spacing after table

def build_doc(filename, persona_name, emails):
    """
    emails: list of dicts with keys: email_num, day, subject, preview, body, cta
    """
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Rho Nutrition")
    title_run.bold = True
    title_run.font.size = Pt(24)
    title_p.paragraph_format.space_after = Pt(4)

    # Subtitle line 1
    sub1 = doc.add_paragraph()
    sub1_run = sub1.add_run(f"Post Purchase Flow Copy | {persona_name}")
    sub1_run.font.size = Pt(13)
    sub1_run.bold = True
    sub1.paragraph_format.space_after = Pt(2)

    # Subtitle line 2
    sub2 = doc.add_paragraph()
    sub2_run = sub2.add_run("Emails 1, 2 + 3  (Day 3 / Day 8 / Day 21)")
    sub2_run.font.size = Pt(11)
    sub2.paragraph_format.space_after = Pt(2)

    # Subtitle line 3
    sub3 = doc.add_paragraph()
    sub3_run = sub3.add_run("DigitsUp x Rho Nutrition | Outersignal Project")
    sub3_run.font.size = Pt(10)
    sub3_run.italic = True
    sub3.paragraph_format.space_after = Pt(16)

    # Emails
    for idx, em in enumerate(emails):
        if idx > 0:
            # Horizontal rule via paragraph border
            hr_p = doc.add_paragraph()
            hr_p.paragraph_format.space_before = Pt(12)
            hr_p.paragraph_format.space_after = Pt(12)
            pPr = hr_p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            pPr.append(pBdr)

        add_email_table(
            doc,
            em['email_num'],
            em['day'],
            persona_name,
            em['subject'],
            em['preview'],
            em['body'],
            em['cta'],
        )

    path = os.path.join(OUTPUT_DIR, filename)
    doc.save(path)
    return path


# ── PERSONA 1: HIGH-PERFORMANCE EXECUTIVE ───────────────────────────────────

P1_E1_BODY = """You made a smart investment. Here is how to protect it.

One serving. Every morning. That is the protocol.

Rho is formulated with liposomal delivery — meaning the nutrients are encapsulated in phospholipid layers that mirror your cell membranes. They absorb faster and more completely than standard capsules. No wasted product. No guessing.

Take it in the morning, on its own or with your existing stack. Consistency from day one is what drives results.

Most people who invest in their performance skip the follow-through. You are already ahead of that."""

P1_E2_BODY = """You are one week in. Here is what is happening at the cellular level.

NAD+ is a coenzyme your cells require for three critical functions: DNA repair, mitochondrial energy production, and cognitive function. NAD+ levels decline with age and under sustained cognitive load.

Rho delivers NAD+ via liposomal encapsulation. Standard capsules lose a significant percentage of potency in the GI tract before reaching target cells. Liposomal delivery bypasses that degradation and increases cellular uptake.

The results from NAD+ are not acute. They compound. Week one establishes the baseline. The variable that determines outcome is consistency.

You are ahead of most people who try to optimize this. Stay on the protocol."""

P1_E3_BODY = """Three weeks is baseline. Not destination.

The cellular work done in weeks one through three does not hold indefinitely without continued support. A gap in your protocol means that work begins to reverse.

NAD+ levels do not stay elevated on their own. Consistent daily dosing is the mechanism.

At less than the cost of a business lunch, your daily protocol maintains the cellular function you have been building.

[NON-SUBSCRIBER VERSION]
Subscription locks in your supply and your price. Skip or cancel anytime — this is a protocol decision, not a commitment. The cost-per-day math is straightforward.

[SUBSCRIBER VERSION]
Your supply is secured. The compounding continues uninterrupted. Manage your subscription anytime from your account."""

P1_EMAILS = [
    dict(email_num=1, day=3, subject="You made a smart call. Here is how to get the most from it.",
         preview="One serving, every morning — here is exactly how to run the protocol.",
         body=P1_E1_BODY, cta="Add to Your Morning Stack"),
    dict(email_num=2, day=8, subject="What NAD+ is doing right now",
         preview="One week in — the cellular mechanism is active, and here is what that means.",
         body=P1_E2_BODY, cta="See the Research"),
    dict(email_num=3, day=21, subject="The protocol only works if you do not stop",
         preview="Three weeks is the baseline, not the finish line — here is what comes next.",
         body=P1_E3_BODY, cta="Secure Your Next Supply"),
]


# ── PERSONA 2: CREATIVE ENTREPRENEUR ────────────────────────────────────────

P2_E1_BODY = """Your mornings deserve something intentional.

Rho is a single daily serving. Take it in the morning — with or without food. Liquid format means it fits easily into whatever your morning already looks like.

There is nothing in here that does not need to be here. Clean ingredients. No fillers. No synthetics. That is a decision the brand made on purpose.

Liposomal delivery means the nutrients actually reach your cells — absorbed more completely than a standard capsule because of the way they are encapsulated.

One serving. One ritual. Every morning. That is all this takes."""

P2_E2_BODY = """What you put in your body is a reflection of what you value.

After one week, here is what is inside Rho and why it matters.

NAD+ supports mental sharpness and cellular energy — the kind of clarity that compounds quietly over time. Glutathione is your body's primary cellular defense compound, helping neutralize oxidative stress that builds from long creative hours and inconsistent rhythms.

What is not in Rho matters as much as what is. No artificial flavors. No synthetic fillers. No unnecessary additives. The label says exactly what it is.

This is not a mass-market wellness product. It was built to meet a different standard.

Keep going. You are one week in."""

P2_E3_BODY = """The ritual is working. Protect it.

Three weeks of a consistent morning practice creates something real. Something worth keeping.

[NON-SUBSCRIBER VERSION]
Subscription means your ritual runs on autopilot. Your supply arrives on schedule. You never run out. You never think about it. Skip or cancel anytime — this is about continuity, not obligation.

[SUBSCRIBER VERSION]
Your ritual continues effortlessly. Everything is already taken care of. You can manage or adjust from your account whenever you like."""

P2_EMAILS = [
    dict(email_num=1, day=3, subject="Your new ritual is ready",
         preview="One clean serving, every morning — here is how to make it yours.",
         body=P2_E1_BODY, cta="Make It Your Morning Ritual"),
    dict(email_num=2, day=8, subject="Clean ingredients for a clear mind",
         preview="One week in -- here is what is inside Rho and why it was built this way.",
         body=P2_E2_BODY, cta="See What Is Inside"),
    dict(email_num=3, day=21, subject="Never run out of what is working",
         preview="Three weeks of consistent mornings is worth protecting -- here is how.",
         body=P2_E3_BODY, cta="Never Run Out"),
]


# ── PERSONA 3: COMMUNITY EDUCATOR AND PARENT ────────────────────────────────

P3_E1_BODY = """You made a good choice. For yourself.

You are here for everyone else most of the time. This one is just for you.

Here is how to get started: one serving, every morning. Take it with or without food. That is all the routine asks.

Rho was built to be simple and clean. No harsh additives. No unnecessary ingredients. The founder started the brand because he believed people deserved a supplement that was actually safe and honest about what was inside.

Liposomal delivery means the nutrients absorb more completely than standard supplements — your body gets more of what you are paying for.

One serving. Every morning. You have got this."""

P3_E2_BODY = """You move through environments full of people. Schools. Community spaces. Home.

High-contact days take a toll on immune resilience, especially when you are the one holding everything together.

Glutathione is your body's master antioxidant — it helps neutralize cellular damage that accumulates from stress and exposure. Rho pairs it with Vitamin C for layered immune support you can actually count on.

These ingredients are clean. No fillers, no harmful additives, nothing you would not want to explain. The brand was founded on that commitment.

One week in, you are building something. The support compounds when you stay consistent.

Keep going."""

P3_E3_BODY = """Three weeks of showing up for yourself. That is real progress.

Immune support does not switch on overnight — it builds layer by layer with consistent daily use. You have already done the hard part.

The people around you benefit when you stay consistent. Teachers, caregivers, parents — your resilience ripples outward.

[NON-SUBSCRIBER VERSION]
Subscription keeps your supply reliable, so you never run out when it matters most. Skip or cancel anytime. This is about the people who count on you having what they need.

[SUBSCRIBER VERSION]
You are all set. Your supply is taken care of. A warm thank you for staying consistent -- manage your subscription from your account anytime."""

P3_EMAILS = [
    dict(email_num=1, day=3, subject="You made a good choice for yourself",
         preview="Here is how to start your Rho routine -- one simple step, every morning.",
         body=P3_E1_BODY, cta="Start Your Routine Today"),
    dict(email_num=2, day=8, subject="Immune support built for your world",
         preview="One week in -- here are the ingredients doing the work and why they matter.",
         body=P3_E2_BODY, cta="Learn What Is Inside"),
    dict(email_num=3, day=21, subject="Stay consistent for the people who count on you",
         preview="Three weeks in is real progress -- here is how to keep the momentum going.",
         body=P3_E3_BODY, cta="Keep Your Supply Going"),
]


# ── PERSONA 4: CLINICAL WELLNESS ADVOCATE ───────────────────────────────────

P4_E1_BODY = """You already understood the formulation before you ordered. That is the right starting point.

For protocol consistency: one dose, taken in the morning. The absorption window for liposomal compounds is optimized when taken at the start of the day, prior to significant food intake.

Rho uses phospholipid bilayer encapsulation to protect active compounds through the GI tract and facilitate direct cellular absorption. This is not a marketing distinction -- it is a meaningful pharmacokinetic difference that impacts bioavailability.

Consistency from day one is the clinical variable that matters most. The mechanism requires sustained plasma levels to drive measurable cellular outcomes.

Review the full formulation below."""

P4_E2_BODY = """Standard oral supplements face significant bioavailability challenges.

Conventional capsule delivery exposes active compounds to gastric acid degradation and first-pass hepatic metabolism, substantially reducing the concentration that reaches target cells.

Liposomal encapsulation via phospholipid bilayers allows hepatic bypass and direct integration with cell membranes, increasing cellular uptake of key compounds including NAD+, Glutathione, and Vitamin C.

NAD+ at the cellular level supports sirtuins activation, mitochondrial biogenesis, and poly(ADP-ribose) polymerase-mediated DNA repair. Glutathione functions as the primary endogenous antioxidant, with systemic depletion implicated in accelerated cellular aging and immune dysregulation.

Cellular saturation requires consistent daily dosing. One week establishes baseline plasma levels. The clinical data builds from here.

Continue the protocol."""

P4_E3_BODY = """Clinical results require clinical consistency.

Liposomal compounds maintain their efficacy through sustained dosing that keeps cellular saturation above the threshold required for measurable function. A gap in the protocol means re-establishing baseline from the beginning -- not continuing from where you left off.

This is the same framework you apply when recommending supplement protocols to patients. The principle holds for your own regimen.

[NON-SUBSCRIBER VERSION]
Subscription eliminates protocol interruptions. Skip or pause anytime -- this is a clinical continuity decision. Consistent cellular saturation requires uninterrupted supply.

[SUBSCRIBER VERSION]
Protocol maintained. Cellular saturation is uninterrupted. Manage your subscription from your account."""

P4_EMAILS = [
    dict(email_num=1, day=3, subject="Your liposomal protocol is active",
         preview="Dosing timing, absorption mechanism, and what to expect from consistent use.",
         body=P4_E1_BODY, cta="View the Full Formulation"),
    dict(email_num=2, day=8, subject="The bioavailability difference, by the numbers",
         preview="One week in -- here is the mechanism driving cellular uptake and why it matters.",
         body=P4_E2_BODY, cta="Read the Research"),
    dict(email_num=3, day=21, subject="Clinical consistency drives clinical results",
         preview="Three weeks in -- here is why maintaining the regimen without interruption matters.",
         body=P4_E3_BODY, cta="Maintain Your Protocol"),
]


# ── PERSONA 5: ACTIVE PERFORMANCE OPTIMIZER ─────────────────────────────────

P5_E1_BODY = """Recovery is not what happens after training. It is the system you build before the next session.

One serving, every morning. Stack it in before your day starts. Consistent daily use is what makes the difference between supplementing and actually recovering.

Rho includes Curcumin and Resveratrol, two compounds with direct anti-inflammatory applications for athletes. Liposomal delivery means those compounds reach the tissue where the work is happening -- not just your GI tract.

Day one is not early. Day one is the only time that matters. Get it in the stack now."""

P5_E2_BODY = """One week in. Here is exactly what is happening.

Curcumin inhibits the NF-kB inflammatory pathway -- one of the primary drivers of exercise-induced inflammation and delayed-onset muscle soreness. Standard Curcumin has poor bioavailability. Liposomal Curcumin reaches the tissue.

Resveratrol supports mitochondrial function and reduces oxidative stress generated by high-output training. The combination targets inflammation and cellular energy production simultaneously.

Week one is the beginning of compounding recovery, not the peak of it. The compounds build in your system. By week four, the difference becomes measurable -- in recovery time, in how you feel between sessions, in output.

You are one week into something that gets better. Stay consistent."""

P5_E3_BODY = """Three weeks is the foundation. Not the finish line.

The anti-inflammatory compounds in your stack are building. Curcumin and Resveratrol require consistent presence in the system to maintain their effect on the inflammatory pathways that matter to your training.

Week four is where compounding recovery becomes measurable. A gap now means starting the process over.

[NON-SUBSCRIBER VERSION]
Curcumin and Resveratrol leave your system when you stop. The progress you have built resets. Subscription keeps you in the stack without having to think about it. Skip or cancel anytime.

[SUBSCRIBER VERSION]
Locked in. Compounding continues. Your next supply is handled -- manage from your account anytime."""

P5_EMAILS = [
    dict(email_num=1, day=3, subject="Recovery starts with day one",
         preview="Here is how to stack Rho into your routine and what it is doing from the first dose.",
         body=P5_E1_BODY, cta="Add It to Your Stack"),
    dict(email_num=2, day=8, subject="What Curcumin is doing inside your joints",
         preview="One week in -- here is the anti-inflammatory science behind your recovery stack.",
         body=P5_E2_BODY, cta="See the Science Behind Recovery"),
    dict(email_num=3, day=21, subject="Three weeks in. The compounding starts now.",
         preview="The foundation is built -- here is what week four and beyond looks like.",
         body=P5_E3_BODY, cta="Stay in the Stack"),
]


# ── PERSONA 6: SKILLED TRADE AND INDUSTRY LEADER ────────────────────────────

P6_E1_BODY = """Here is exactly how to use it.

One serving. Every morning. That is the whole routine.

Here is why the format matters: most supplements go through your stomach and a significant portion never makes it into your bloodstream. Rho uses liposomal delivery, which means the nutrients are wrapped in a protective layer that gets them absorbed properly. Your money is actually absorbing.

Joint support takes weeks to build up. That is normal. Do not expect to feel it on day three. Stick with it.

Morning. Every day. That is it."""

P6_E2_BODY = """Most supplements do not work the way the label implies.

Standard capsules pass through the digestive system and lose a large percentage of their active compounds before they reach your bloodstream. That is not a Rho problem -- that is a delivery method problem.

Liposomal technology wraps the nutrients in a layer that protects them through your GI tract and gets them absorbed into your bloodstream directly. That is the difference you paid for.

Curcumin for joint support takes two to four weeks to build up. You are not behind schedule. You are on it. The support is accumulating in your system right now.

Keep going."""

P6_E3_BODY = """Three weeks in. Joint support is building.

Do not let it run out and start over from zero.

[NON-SUBSCRIBER VERSION]
Subscription locks in your supply. Skip or cancel anytime. At less than a few dollars a day, you are not paying more -- you are just not stopping what is working.

[SUBSCRIBER VERSION]
You are sorted. Supply is handled. Keep going -- manage from your account if you ever need to adjust."""

P6_EMAILS = [
    dict(email_num=1, day=3, subject="You bought it. Here is exactly how to use it.",
         preview="One serving, every morning -- and why the delivery method makes it worth it.",
         body=P6_E1_BODY, cta="Start Today"),
    dict(email_num=2, day=8, subject="Why this works when others do not",
         preview="One week in -- here is the plain-language explanation of what makes Rho different.",
         body=P6_E2_BODY, cta="See Why It Is Different"),
    dict(email_num=3, day=21, subject="Three weeks in. Lock it in.",
         preview="Joint support is building -- do not let it run out and reset your progress.",
         body=P6_E3_BODY, cta="Lock It In"),
]


# ── BUILD ALL 6 ──────────────────────────────────────────────────────────────

DOCS = [
    ("Rho Nutrition Copy - P1 High-Performance Executive.docx",
     "P1 | The High-Performance Executive", P1_EMAILS),
    ("Rho Nutrition Copy - P2 Creative Entrepreneur.docx",
     "P2 | The Creative Entrepreneur", P2_EMAILS),
    ("Rho Nutrition Copy - P3 Community Educator and Parent.docx",
     "P3 | The Community Educator and Parent", P3_EMAILS),
    ("Rho Nutrition Copy - P4 Clinical Wellness Advocate.docx",
     "P4 | The Clinical Wellness Advocate", P4_EMAILS),
    ("Rho Nutrition Copy - P5 Active Performance Optimizer.docx",
     "P5 | The Active Performance Optimizer", P5_EMAILS),
    ("Rho Nutrition Copy - P6 Skilled Trade and Industry Leader.docx",
     "P6 | The Skilled Trade and Industry Leader", P6_EMAILS),
]

paths = []
for filename, persona_name, emails in DOCS:
    p = build_doc(filename, persona_name, emails)
    paths.append(p)
    print(f"Saved: {p}")

print("\nAll 6 documents complete.")
