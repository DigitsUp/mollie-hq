from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p

def add_label(doc, label, value):
    p = doc.add_paragraph()
    run = p.add_run(label + ": ")
    run.bold = True
    run.font.size = Pt(11)
    p.add_run(value).font.size = Pt(11)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    return p

def add_section_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CCCCCC')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_brief(doc, persona_name, persona_label, email_obj, subject_lines, preview_text, primary_message, content_blocks, cta_direction, tone_notes):
    add_heading(doc, f"{persona_label} — {persona_name}", level=2)

    add_section_title(doc, "Email Objective")
    doc.add_paragraph(email_obj)

    add_section_title(doc, "Subject Line Directions")
    for sl in subject_lines:
        add_bullet(doc, sl)

    add_section_title(doc, "Preview Text Direction")
    doc.add_paragraph(preview_text)

    add_section_title(doc, "Primary Message")
    doc.add_paragraph(primary_message)

    add_section_title(doc, "Content Block Structure")
    for block in content_blocks:
        add_bullet(doc, block)

    add_section_title(doc, "CTA Direction")
    doc.add_paragraph(cta_direction)

    add_section_title(doc, "Tone Notes")
    doc.add_paragraph(tone_notes)

    add_divider(doc)

# ─────────────────────────────────────────────
# TITLE PAGE
# ─────────────────────────────────────────────
doc.add_heading("Rho Nutrition", 0)
p = doc.add_paragraph("Post Purchase Flow — Email Briefs v1")
p.runs[0].font.size = Pt(14)
p = doc.add_paragraph("DigitsUp x Rho Nutrition")
p.runs[0].font.size = Pt(11)
p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
doc.add_paragraph("")

# ─────────────────────────────────────────────
# DAY 3 — EMAIL 1
# ─────────────────────────────────────────────
add_heading(doc, "Day 3 — Email 1: Welcome + How to Use", level=1)
p = doc.add_paragraph("Product has arrived or is arriving. Customer is starting their routine. Primary goal: activation. Get them using the product correctly from day one.")
p.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_paragraph("")

# P1 Executive
add_brief(
    doc,
    persona_name="The High-Performance Executive",
    persona_label="P1",
    email_obj="Reinforce the purchase decision and frame daily usage as a performance protocol. Remove any friction from starting. Position Rho as a high-ROI addition to an already optimized life.",
    subject_lines=[
        "Your optimization protocol starts now",
        "One serving. Every morning. Here is why it matters.",
        "You made a smart call. Here is how to get the most from it.",
    ],
    preview_text="Short, efficient. Mirror their mindset: this fits into your system. Example: 'Takes 30 seconds. Works all day.'",
    primary_message="You made a decision to invest in your performance. Now activate it. One serving of Rho in the morning is all it takes. Liquid liposomal format means faster absorption than capsules — no waiting, no wasted product. Your NAD+ and cognitive support starts on day one.",
    content_blocks=[
        "Hero block: Welcome message — confident, not warm. Acknowledge the purchase as a smart decision.",
        "Usage block: Simple 3-step protocol (when, how, how much). No clutter. Treat it like a systems brief.",
        "What to expect block: Short timeline — Day 1 to Day 30. Frame as compounding ROI, not instant results.",
        "Ingredient callout: NAD+ and one supporting compound. One sentence each. Clinical but accessible.",
        "CTA block: Single action — 'Add to your morning stack' with a link to product page or account.",
    ],
    cta_direction="Direct and action-oriented. 'Add to your morning stack.' No urgency, no discount. Just the next logical step.",
    tone_notes="Boardroom-level confidence. Zero fluff. Short sentences. Think: internal memo, not a wellness newsletter. This person skims — every line must earn its place."
)

# P2 Creative
add_brief(
    doc,
    persona_name="The Creative Entrepreneur",
    persona_label="P2",
    email_obj="Welcome the customer into the Rho ritual. Frame daily usage as an intentional, elevated moment — not a chore. Make starting feel inspiring, not clinical.",
    subject_lines=[
        "Your new ritual is ready",
        "The morning routine just got an upgrade",
        "Clean. Simple. Yours.",
    ],
    preview_text="Warm and sensory. Example: 'A 30-second ritual that sets the tone for everything after.'",
    primary_message="You chose a product that fits the way you live — clean ingredients, thoughtful formulation, nothing unnecessary. The liquid format is intentional: easier to absorb, easier to build into a morning you already care about. This is not just a supplement. It is a small act of taking care of the work you do.",
    content_blocks=[
        "Hero block: Aesthetic welcome. Clean visual reference (light, morning, ritual). Warm tone, not clinical.",
        "Ritual framing block: How to incorporate Rho into a morning routine. Sensory and simple.",
        "Ingredient transparency block: What is in it and what is not. Clean label as a point of pride.",
        "Brand story beat: One line about why Rho was formulated this way. Founder lens if appropriate.",
        "CTA block: 'Make it your morning ritual' — link to product page.",
    ],
    cta_direction="Soft and inviting. 'Make it your morning ritual.' Feels like an invitation, not an instruction.",
    tone_notes="Warm, considered, artisan. Reads like a well-designed brand they already love. Avoid clinical language. Avoid anything that sounds like mass-market wellness. White space is your friend."
)

# P3 Educator/Parent
add_brief(
    doc,
    persona_name="The Community Educator and Parent",
    persona_label="P3",
    email_obj="Make the customer feel good about their decision. Provide clear, reassuring usage instructions. Reinforce that this is a safe, trusted product built for real people with full lives.",
    subject_lines=[
        "You made a good choice for yourself",
        "Here is how to get started with Rho",
        "Simple steps for your new daily routine",
    ],
    preview_text="Warm and affirming. Example: 'You deserve this. Here is how to make it work for your day.'",
    primary_message="Between your students, your family, and everything else you carry — this is one thing that is just for you. Rho is clean, safe, and simple. One serving in the morning. That is all it takes to start supporting your immune system, your energy, and your ability to keep showing up for the people who count on you.",
    content_blocks=[
        "Hero block: Warm welcome. Acknowledge how much they give. This is for them.",
        "Usage block: Simple, friendly instructions. No jargon. 'One serving in the morning with or without food.'",
        "Safety and trust block: Clean ingredients, no harmful additives, founder mission. Two to three short sentences.",
        "What to expect block: Gentle, realistic timeline framed around feeling steady, not dramatic transformation.",
        "CTA block: 'Start your routine today' — simple and encouraging.",
    ],
    cta_direction="Encouraging and low-pressure. 'Start your routine today.' Feels like a nudge from a trusted friend.",
    tone_notes="Human, warm, reassuring. Reads like a message from someone who genuinely cares. No science-heavy language. No performance framing. This person responds to trust and community, not optimization."
)

# P4 Clinical
add_brief(
    doc,
    persona_name="The Clinical Wellness Advocate",
    persona_label="P4",
    email_obj="Confirm the clinical rationale for their purchase. Provide precise dosing and timing guidance. Acknowledge their expertise — they already know why liposomal delivery matters. Confirm they made the right call.",
    subject_lines=[
        "Your liposomal protocol is active",
        "Dosing, timing, and what to expect — your clinical guide",
        "The formulation you chose. Here is how to use it correctly.",
    ],
    preview_text="Clinical and direct. Example: 'Optimal absorption window, full formulation breakdown, and what to watch for.'",
    primary_message="You understand bioavailability. You know why liposomal delivery outperforms standard capsule formulations. Rho uses phospholipid encapsulation to protect active compounds through digestion and deliver them at the cellular level. Take one serving in the morning, ideally fasting or with a light meal for optimal absorption. Glutathione and NAD+ require consistency — clinical outcomes compound over weeks, not days.",
    content_blocks=[
        "Hero block: Peer-level acknowledgment. No hand-holding. 'You already know what liposomal means. Here is how to use Rho correctly.'",
        "Dosing protocol block: Timing, amount, fasting notes, absorption window. Precise and formatted for scanning.",
        "Formulation detail block: Key active compounds (Glutathione, NAD+, Curcumin, Resveratrol), delivery mechanism, why each matters.",
        "Clinical timeline block: What to expect in week 1, week 2, week 4. Framed as cellular — not symptomatic.",
        "CTA block: 'View the full formulation' — link to product page or COA if available.",
    ],
    cta_direction="Evidence-driven. 'View the full formulation.' Invites them to go deeper without overselling.",
    tone_notes="Peer-level. Assume they understand the science. Use correct terminology: bioavailability, liposomal, phospholipid, cellular absorption. Confident and precise. Never dumbed down. Never enthusiastic."
)

# P5 Performance
add_brief(
    doc,
    persona_name="The Active Performance Optimizer",
    persona_label="P5",
    email_obj="Connect Rho directly to their training and recovery routine. Frame day one as the start of a performance system. Give them a clear, actionable protocol they can stack immediately.",
    subject_lines=[
        "Recovery starts with day one",
        "Here is how to stack Rho into your routine",
        "Day one. The protocol begins.",
    ],
    preview_text="Action-oriented. Example: 'Morning timing, anti-inflammatory support, and what you should start noticing by week two.'",
    primary_message="Recovery is not what happens after you train. It is the system you build before the next session. Rho fits into that system. Take it in the morning — Curcumin and Resveratrol work best with consistent daily intake, not just post-workout. By day eight you should notice the difference in how your joints feel during and after training. Start the protocol today.",
    content_blocks=[
        "Hero block: Performance framing. 'Recovery does not start after the session. It starts here.'",
        "Stacking protocol block: When to take it, why morning timing matters for anti-inflammatory compounds. Practical and direct.",
        "Ingredient callout block: Curcumin + Resveratrol — what they do, why bioavailability matters for performance. One paragraph, no fluff.",
        "Timeline block: What to expect in the first week. Name the physical signals — joint comfort, recovery speed. Validate what they will feel.",
        "CTA block: 'Add it to your stack' — link to product page.",
    ],
    cta_direction="Active and peer-level. 'Add it to your stack.' Sounds like advice from a trusted coach, not a brand.",
    tone_notes="Knowledgeable, direct, energetic. Reads like a message from a fellow coach or athlete. No lifestyle language. No ritual language. Results and recovery are the only currencies here."
)

# P6 Trade
add_brief(
    doc,
    persona_name="The Skilled Trade and Industry Leader",
    persona_label="P6",
    email_obj="Get straight to it. Tell them how to take it and what to expect. No story, no ritual, no lifestyle framing. Practical instructions from someone who respects their time.",
    subject_lines=[
        "Here is how to get the most out of it",
        "Simple. Consistent. Effective. Your guide to Rho.",
        "You bought it. Here is exactly how to use it.",
    ],
    preview_text="Blunt and practical. Example: 'One serving a day. Morning is best. Here is what consistent use does.'",
    primary_message="Take one serving in the morning. That is it. Do it every day and the compounds build up in your system. Curcumin and Resveratrol work on joint inflammation over time — not overnight. Liposomal delivery means your body actually absorbs what you paid for, unlike most supplements. Stick with it.",
    content_blocks=[
        "Hero block: No fluff welcome. 'You made a solid decision. Here is how to use it right.'",
        "Usage instructions block: Straight directive. When, how much, with or without food. Bullet format.",
        "Why it works block: Liposomal vs standard — one short paragraph. Frame as: this is why you are not wasting your money.",
        "What to expect block: Realistic and honest. Joint support and energy take a few weeks. No hype.",
        "CTA block: 'Start today' — link to product page.",
    ],
    cta_direction="Blunt and direct. 'Start today.' No decoration.",
    tone_notes="Straight talk. Short sentences. Zero lifestyle or wellness language. This person respects honesty and efficiency. If it sounds like a wellness newsletter, rewrite it."
)

# Save
path = "/Users/minigrill/.openclaw/workspace/projects/outersignal/rho-nutrition/Rho Nutrition Post Purchase Briefs v1.docx"
doc.save(path)
print(f"Saved: {path}")
