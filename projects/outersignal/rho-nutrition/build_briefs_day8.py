from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import copy

def add_persona_brief(doc, email_label, persona_name, heading_name, data):
    # Heading 2
    h = doc.add_heading(heading_name, level=2)

    # Table: 2 cols, 8 rows
    table = doc.add_table(rows=8, cols=2)
    table.style = 'Table Grid'

    rows = table.rows

    # Row 0: header - both cells same
    rows[0].cells[0].text = email_label + '  |  ' + persona_name.upper()
    rows[0].cells[1].text = email_label + '  |  ' + persona_name.upper()

    fields = [
        ('Email Objective', data['objective']),
        ('Subject Line Directions', data['subjects']),
        ('Preview Text Direction', data['preview']),
        ('Primary Message', data['primary']),
        ('Content Block Structure', data['blocks']),
        ('CTA Direction', data['cta']),
        ('Tone Notes', data['tone']),
    ]
    for i, (label, value) in enumerate(fields):
        rows[i+1].cells[0].text = label
        rows[i+1].cells[1].text = value

    doc.add_paragraph('')


# ── DAY 8 ──
d8 = Document()
d8.add_paragraph('Rho Nutrition', style='Title')
d8.add_paragraph('Post Purchase Flow — Email Briefs v1')
d8.add_paragraph('DigitsUp x Rho Nutrition  |  All 6 Personas  |  Day 8: Product Education + Results Timeline')

personas_day8 = [
    {
        'heading': 'P1: The High-Performance Executive',
        'label': 'EMAIL 2 — DAY 8',
        'name': 'The High-Performance Executive',
        'objective': 'Validate the decision with data. One week in, explain what NAD+ is doing at the cellular level. Build belief through science, not testimonials.',
        'subjects': 'What NAD+ is doing right now\nThe cellular results you cannot see yet\nOne week in. Here is what is happening.',
        'preview': 'Direct, data-led. Example: "Cellular energy, cognitive clarity, compounding results. The science behind what you are feeling."',
        'primary': 'At the cellular level, NAD+ is repairing DNA, fueling mitochondria, and supporting cognitive function. The results compound. Explain the mechanism briefly, confirm they are ahead of most people, and reinforce daily consistency as the performance variable.',
        'blocks': 'Hero: Week one validation, data-led. Brief acknowledgment of decision quality.\nSecondary: What NAD+ does at the cellular level. 2-3 sentences, precise. Mitochondrial function, DNA repair, cognitive support.\nSecondary: Liposomal delivery vs capsules. Why absorption matters and what they chose over standard formats.\nTertiary: Compounding timeline. Week 2, 4, 8. Frame as ROI curve. Ends with: consistency is the protocol.',
        'cta': '"Read the Science" or "See the Research" — evidence-driven, no urgency.',
        'tone': 'Data-led, efficient, precise. No enthusiasm. No wellness language. Boardroom voice throughout.',
    },
    {
        'heading': 'P2: The Creative Entrepreneur',
        'label': 'EMAIL 2 — DAY 8',
        'name': 'The Creative Entrepreneur',
        'objective': 'Connect clean ingredients to mental clarity and creative output. One week in — validate what they may be noticing. Build ingredient transparency as a brand value.',
        'subjects': 'Clean ingredients for a clear mind\nOne week in. What are you noticing?\nThe label that says exactly what it is.',
        'preview': 'Warm and reflective. Example: "What is inside, why it matters, and what a week of clean nutrition can quietly shift."',
        'primary': 'Clean label. No fillers, no synthetics, no surprises. Walk through the key ingredients in a way that feels like brand storytelling, not a clinical breakdown. NAD+ for mental sharpness. Glutathione for cellular clarity. It fits the way they live.',
        'blocks': 'Hero: One week check-in. Warm, curious. "What are you noticing?" Gentle, not clinical.\nSecondary: Ingredient transparency. Key compounds named, one line each, lifestyle framing. NAD+, Glutathione, Vitamin C.\nSecondary: Clean label philosophy. What is not in Rho is as important as what is. No fillers, no synthetics.\nTertiary: Connection to creative work. Mental clarity, sustained energy, no crash. Short and reflective.',
        'cta': '"See What Is Inside" — feels like an invitation, not an instruction.',
        'tone': 'Considered, honest, artisan. Reads like a brand they already love. No clinical language. No mass-market wellness.',
    },
    {
        'heading': 'P3: The Community Educator and Parent',
        'label': 'EMAIL 2 — DAY 8',
        'name': 'The Community Educator and Parent',
        'objective': 'Reinforce immune support and safety. One week in — acknowledge the high-demand environment they work in. Build trust through ingredient transparency and founder mission.',
        'subjects': 'Immune support built for your world\nOne week in. Keep going.\nThe ingredients you can trust.',
        'preview': 'Warm and affirming. Example: "You are around a lot of people every day. Here is what Rho is doing on your behalf."',
        'primary': 'High-contact environments demand stronger immune support. Glutathione is the body\'s master antioxidant. Vitamin C supports immune response. This is not about performance. It is about staying well so you can keep showing up. Founder mission reinforces trust.',
        'blocks': 'Hero: One week warm check-in. "You are doing it. Keep going." Acknowledges effort, not performance.\nSecondary: Immune support callout. Glutathione + Vitamin C, explained simply and warmly. No jargon.\nSecondary: Clean and safe ingredients. No harmful additives. Founder story in 2 sentences. Trust-building, human.\nTertiary: Gentle encouragement to stick with the routine for cumulative benefit. Framed around showing up for others.',
        'cta': '"Learn What Is Inside" — warm, reassuring, low-pressure.',
        'tone': 'Community, care, trust. No performance framing. Simple language throughout. Human first.',
    },
    {
        'heading': 'P4: The Clinical Wellness Advocate',
        'label': 'EMAIL 2 — DAY 8',
        'name': 'The Clinical Wellness Advocate',
        'objective': 'Go deep on the bioavailability science. Peer-level education on liposomal delivery vs standard oral supplements. Validate their recommendation confidence.',
        'subjects': 'The bioavailability difference, by the numbers\nLiposomal delivery: what the research shows\nOne week in. The mechanism is working.',
        'preview': 'Clinical and precise. Example: "Phospholipid encapsulation, absorption rates, and why liposomal is the standard your patients deserve."',
        'primary': 'Standard oral supplements lose significant potency before reaching target cells. Liposomal encapsulation using phospholipid bilayers protects active compounds through the GI tract and delivers them at the cellular level. The data is clear. This is the formulation worth recommending.',
        'blocks': 'Hero: Peer acknowledgment. "One week in. The mechanism is active." Clinical, no congratulations.\nSecondary: Liposomal vs standard. Absorption comparison, mechanism explained with correct terminology: phospholipid bilayer, hepatic bypass, cellular uptake.\nSecondary: Key compounds and clinical significance. Glutathione (antioxidant capacity, detoxification), NAD+ (mitochondrial function, DNA repair), Curcumin (NF-kB pathway, anti-inflammatory).\nTertiary: Patient recommendation framing. What to tell patients. Why this formulation supports clinical outcomes. Confidence-building for the clinician.',
        'cta': '"Read the Research" — evidence-driven, peer-appropriate.',
        'tone': 'Peer-level throughout. Correct terminology only. No simplification. No enthusiasm. Precise and confident.',
    },
    {
        'heading': 'P5: The Active Performance Optimizer',
        'label': 'EMAIL 2 — DAY 8',
        'name': 'The Active Performance Optimizer',
        'objective': 'Validate the physical signals they may be noticing. Explain what Curcumin and Resveratrol are doing inside joints and muscles. Build belief through mechanism, not testimonial.',
        'subjects': 'What Curcumin is doing inside your joints\nOne week in. Here is what is happening.\nThe anti-inflammatory science behind your recovery.',
        'preview': 'Active and direct. Example: "Curcumin, Resveratrol, inflammation pathways, and what consistent use means for your next session."',
        'primary': 'Curcumin inhibits inflammatory pathways at the molecular level. Resveratrol supports mitochondrial function and reduces oxidative stress. In a liposomal format, these compounds actually reach the tissues where you need them. One week is not the peak. It is the start of compounding recovery.',
        'blocks': 'Hero: One week in validation. "You might be noticing something. Here is the science behind it." Direct, no fluff.\nSecondary: Curcumin mechanism. NF-kB pathway, joint inflammation, why liposomal matters for bioavailability. Performance-first framing.\nSecondary: Resveratrol. Mitochondrial function, oxidative stress, recovery speed. Active, coach-like language.\nTertiary: Compounding recovery timeline. Week 2, 4, 8. What to expect physically. Validates what they are already feeling.',
        'cta': '"See the Science Behind Recovery" — active and peer-level.',
        'tone': 'Knowledgeable, energetic, direct. Results and recovery are the only currencies. No lifestyle language.',
    },
    {
        'heading': 'P6: The Skilled Trade and Industry Leader',
        'label': 'EMAIL 2 — DAY 8',
        'name': 'The Skilled Trade and Industry Leader',
        'objective': 'Explain why this works when other supplements did not. Liposomal vs standard in plain money language. Validate joint health support for physical demands. Keep it short and credible.',
        'subjects': 'Why this works when others do not\nOne week in. Here is what to expect next.\nThe difference you paid for, explained plainly.',
        'preview': 'Blunt and practical. Example: "Most supplements never make it to where you need them. Here is why Rho does."',
        'primary': 'Most supplements pass through without absorbing. Liposomal means the nutrients actually get into your bloodstream. Joint support from Curcumin takes a few weeks to build. That is normal. Keep going. You are not wasting your money.',
        'blocks': 'Hero: Straight talk. "One week in. Here is what is happening." Short. No buildup.\nSecondary: Liposomal explained simply. One short paragraph. Framed as: your money is actually working. No scientific jargon.\nSecondary: Joint health timeline. Curcumin takes 2-4 weeks to feel it. Honest, no hype. Sets expectations accurately.\nTertiary: Simple encouragement. Stick with it. That is the whole message. One or two sentences maximum.',
        'cta': '"See Why It Is Different" — plain, direct.',
        'tone': 'Straight, credible, practical. Zero lifestyle language. Short sentences. No wellness newsletter energy.',
    },
]

for p in personas_day8:
    add_persona_brief(d8, p['label'], p['name'], p['heading'], p)

out8 = '/Users/minigrill/.openclaw/workspace/projects/outersignal/rho-nutrition/Rho Nutrition Post Purchase Briefs Day 8 v1.docx'
d8.save(out8)
print(f'Saved: {out8}')


# ── DAY 21 ──
d21 = Document()
d21.add_paragraph('Rho Nutrition', style='Title')
d21.add_paragraph('Post Purchase Flow — Email Briefs v1')
d21.add_paragraph('DigitsUp x Rho Nutrition  |  All 6 Personas  |  Day 21: Reorder + Subscription Conversion')

personas_day21 = [
    {
        'heading': 'P1: The High-Performance Executive',
        'label': 'EMAIL 3 — DAY 21',
        'name': 'The High-Performance Executive',
        'objective': 'Frame consistency as a protocol, not a habit. Cost-per-day logic. Prevent a gap in the stack.',
        'subjects': 'Do not let your results reset\nThe protocol only works if you do not stop\nYour next supply. Before you need it.',
        'preview': 'Logical and outcome-focused. Example: "Three weeks of cellular optimization. The ROI compounds only if the protocol continues."',
        'primary': 'Three weeks of consistent NAD+ and liposomal nutrition is the baseline, not the destination. The results compound with time. Stopping resets the cellular work. Cost-per-day framing: compare to a single business lunch. The protocol is the competitive edge. Do not interrupt it.',
        'blocks': 'Hero: Three weeks in. Brief data-driven acknowledgment of progress. Tone: boardroom, not celebratory.\n\nDYNAMIC BLOCK (Secondary):\nSUBSCRIBER VERSION: Acknowledge they are locked in. Reinforce the decision with one line about compounding results. CTA to account dashboard.\nNON-SUBSCRIBER VERSION: Cost-per-day math. Compare to cost of a single business lunch. Never run out. Skip or cancel anytime. CTA to subscribe.\n\nTertiary: One short paragraph on what happens when you stop. NAD+ depletion timeline. Urgency through logic, not hype.',
        'cta': '"Secure Your Next Supply" (non-subscriber) / "You Are Locked In" (subscriber)',
        'tone': 'Logical, outcome-focused. Boardroom throughout. Zero sentiment.',
    },
    {
        'heading': 'P2: The Creative Entrepreneur',
        'label': 'EMAIL 3 — DAY 21',
        'name': 'The Creative Entrepreneur',
        'objective': 'Protect the ritual. Frame subscription as never having to think about it again. Identity-affirming.',
        'subjects': 'Keep the ritual going\nYour routine is worth protecting\nNever run out of what is working',
        'preview': 'Gentle and identity-affirming. Example: "Three weeks of something that is just yours. Here is how to keep it."',
        'primary': 'The ritual is working. Do not let it lapse because of logistics. Subscription is the simplest version of self-care: the thing that works shows up when it should. Frame it as removing friction from something worth protecting.',
        'blocks': 'Hero: Three weeks in. Warm acknowledgment. "This has become part of how you start your day."\n\nDYNAMIC BLOCK (Secondary):\nSUBSCRIBER VERSION: Warm affirmation. You are already set. The ritual continues effortlessly. Reinforce the decision gently. CTA to account.\nNON-SUBSCRIBER VERSION: Subscription as ritual protection. Never run out framing. Skip or cancel anytime note. Single clear CTA.\n\nTertiary: One short line connecting consistent clean nutrition to sustained creative clarity. Gentle close.',
        'cta': '"Never Run Out" (non-subscriber) / "Your Ritual Continues" (subscriber)',
        'tone': 'Warm, gentle, identity-affirming. No urgency. No discount language.',
    },
    {
        'heading': 'P3: The Community Educator and Parent',
        'label': 'EMAIL 3 — DAY 21',
        'name': 'The Community Educator and Parent',
        'objective': 'Reliability for the people who count on them. Frame consistency as showing up for others. Low-pressure.',
        'subjects': 'Stay consistent for the people who count on you\nThree weeks in. Keep going.\nYour routine is working. Do not let it lapse.',
        'preview': 'Warm and community-centered. Example: "You show up for everyone else. This is one way to show up for yourself, consistently."',
        'primary': 'Three weeks is real progress. Immune support and sustained energy do not stop building. They compound with time. The people around you benefit when you stay consistent. Keep the supply going. This is not a luxury. It is maintenance.',
        'blocks': 'Hero: Warm three-week acknowledgment. "You have been consistent. That matters."\n\nDYNAMIC BLOCK (Secondary):\nSUBSCRIBER VERSION: Warm affirmation. Reinforce they are already set. Thank them. CTA to account.\nNON-SUBSCRIBER VERSION: Never run out framing. Skip or cancel note. Framed around reliability for family and community. CTA to subscribe.\n\nTertiary: Gentle closing. "You deserve this too." Short and warm. One or two sentences.',
        'cta': '"Keep Your Supply Going" (non-subscriber) / "You Are All Set" (subscriber)',
        'tone': 'Nurturing, community-first. No performance language. No urgency. Simple and warm throughout.',
    },
    {
        'heading': 'P4: The Clinical Wellness Advocate',
        'label': 'EMAIL 3 — DAY 21',
        'name': 'The Clinical Wellness Advocate',
        'objective': 'Clinical consistency as protocol. Uninterrupted supplementation produces measurable outcomes. Professional, outcome-driven.',
        'subjects': 'Clinical consistency drives clinical results\nDo not interrupt the protocol\nThree weeks in. Maintain the regimen.',
        'preview': 'Clinical and direct. Example: "Interrupted supplementation resets cellular baseline. Maintaining the protocol is the clinical decision."',
        'primary': 'Liposomal supplementation requires consistent dosing to maintain cellular saturation. A gap in the protocol means reestablishing baseline. For patients you may recommend this to, the message is the same: uninterrupted use produces measurable outcomes. Maintain the protocol.',
        'blocks': 'Hero: Clinical three-week milestone. Brief acknowledgment. Tone: peer, not congratulatory.\n\nDYNAMIC BLOCK (Secondary):\nSUBSCRIBER VERSION: Acknowledge protocol maintained. Confirm cellular work is uninterrupted. CTA to account dashboard.\nNON-SUBSCRIBER VERSION: Interrupted supplementation data point. Cost of stopping and restarting. Skip or cancel framing. CTA to subscribe.\n\nTertiary: Patient recommendation framing. Consistent liposomal use over 90 days is where outcomes become measurable. Builds recommendation confidence.',
        'cta': '"Maintain Your Protocol" (non-subscriber) / "Protocol Maintained" (subscriber)',
        'tone': 'Professional, clinical, outcome-driven. No warmth. No lifestyle framing. Peer-level throughout.',
    },
    {
        'heading': 'P5: The Active Performance Optimizer',
        'label': 'EMAIL 3 — DAY 21',
        'name': 'The Active Performance Optimizer',
        'objective': 'Consistency is compounding recovery. Do not lose the progress built over three weeks. Coach-like energy.',
        'subjects': 'Do not lose the progress you have built\nThree weeks in. The compounding starts now.\nStay in the stack.',
        'preview': 'Motivating and direct. Example: "Three weeks of Curcumin and Resveratrol building in your system. Week four is where it compounds."',
        'primary': 'Three weeks of consistent anti-inflammatory support is not the peak. It is the foundation. The compounds are building in your system. Week four and beyond is where compounding recovery becomes measurable. A gap means starting over. Stay in the stack.',
        'blocks': 'Hero: Three weeks. Performance validation. "You are building something. Do not stop now."\n\nDYNAMIC BLOCK (Secondary):\nSUBSCRIBER VERSION: You are locked in. Compounding continues. Brief acknowledgment. CTA to account.\nNON-SUBSCRIBER VERSION: What stops when you stop. Curcumin and Resveratrol leave the system in 1-2 weeks. Skip or cancel note. CTA to subscribe.\n\nTertiary: Week 4-8 timeline. What to expect next. Performance framing. Bridges to long-term commitment.',
        'cta': '"Stay in the Stack" (non-subscriber) / "Locked In. Keep Going." (subscriber)',
        'tone': 'Coach-like, motivating, direct. Performance language only. No lifestyle or ritual framing.',
    },
    {
        'heading': 'P6: The Skilled Trade and Industry Leader',
        'label': 'EMAIL 3 — DAY 21',
        'name': 'The Skilled Trade and Industry Leader',
        'objective': 'Running low? Simple. Do not let it run out. No restart needed. Lock it in.',
        'subjects': 'Running low? Do not let it run out.\nThree weeks in. Lock it in.\nSimple. Keep it going.',
        'preview': 'Blunt and practical. Example: "You are three weeks in. Do not start over."',
        'primary': 'You are three weeks in. The joint support is building. Do not let it run out and start over. Lock in a subscription and forget about it. Skip or cancel anytime. Just do not let it reset.',
        'blocks': 'Hero: Three weeks. Short, direct. "Good. Keep going."\n\nDYNAMIC BLOCK (Secondary):\nSUBSCRIBER VERSION: You are sorted. Keep going. CTA to account. Two sentences maximum.\nNON-SUBSCRIBER VERSION: Do not start over framing. Subscription means locked in. Skip or cancel note. Cost-per-day in plain language. CTA to subscribe.\n\nTertiary: One line only. No extra content needed for this persona.',
        'cta': '"Lock It In" (non-subscriber) / "You Are Sorted" (subscriber)',
        'tone': 'Blunt, short, zero decoration. Straight talk. If it sounds like a wellness newsletter, rewrite it.',
    },
]

for p in personas_day21:
    add_persona_brief(d21, p['label'], p['name'], p['heading'], p)

out21 = '/Users/minigrill/.openclaw/workspace/projects/outersignal/rho-nutrition/Rho Nutrition Post Purchase Briefs Day 21 v1.docx'
d21.save(out21)
print(f'Saved: {out21}')
