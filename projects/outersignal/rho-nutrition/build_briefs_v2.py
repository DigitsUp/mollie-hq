from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '999999')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_section_header(table, title, color='1A1A1A'):
    row = table.add_row()
    row.cells[0].merge(row.cells[1])
    cell = row.cells[0]
    set_cell_bg(cell, color)
    set_cell_border(cell)
    p = cell.paragraphs[0]
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

def add_block_header(table, title, block_type):
    row = table.add_row()
    row.cells[0].merge(row.cells[1])
    cell = row.cells[0]
    set_cell_bg(cell, 'E8F0E8')
    set_cell_border(cell)
    p = cell.paragraphs[0]
    run = p.add_run(f"{title}  ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    tag = p.add_run(f"[{block_type}]")
    tag.font.size = Pt(9)
    tag.font.italic = True
    tag.font.color.rgb = RGBColor(0x55, 0x88, 0x55)
    tag.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

def add_row(table, label, value, is_note=False):
    row = table.add_row()
    lc = row.cells[0]
    set_cell_border(lc)
    set_cell_bg(lc, 'F5F5F5')
    p = lc.paragraphs[0]
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

    vc = row.cells[1]
    set_cell_border(vc)
    lines = value.split('\n')
    first = True
    for line in lines:
        if first:
            p = vc.paragraphs[0]
            first = False
        else:
            p = vc.add_paragraph()
        run = p.add_run(line)
        run.font.size = Pt(9 if is_note else 10)
        run.font.italic = is_note
        run.font.name = 'Calibri'
        if is_note:
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        p.paragraph_format.space_after = Pt(2)

def add_bullets(table, label, items):
    row = table.add_row()
    lc = row.cells[0]
    set_cell_border(lc)
    set_cell_bg(lc, 'F5F5F5')
    p = lc.paragraphs[0]
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'

    vc = row.cells[1]
    set_cell_border(vc)
    first = True
    for item in items:
        if first:
            p = vc.paragraphs[0]
            first = False
        else:
            p = vc.add_paragraph()
        p.style = 'List Bullet'
        run = p.add_run(item)
        run.font.size = Pt(10)
        run.font.name = 'Calibri'
        p.paragraph_format.space_after = Pt(2)

def add_persona_brief(doc, persona_code, persona_name, email_num, send_day,
                      objective, subject_lines, preview_text_direction,
                      primary_message, content_blocks, cta_direction, tone_notes):

    h = doc.add_heading(f"{persona_code}: {persona_name}", level=2)
    h.paragraph_format.space_before = Pt(20)
    h.paragraph_format.space_after = Pt(8)

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(4.8)

    add_section_header(table, f"Email {email_num} — Day {send_day}  |  {persona_name}")
    add_row(table, "Email Objective", objective)
    add_bullets(table, "Subject Line Directions", subject_lines)
    add_row(table, "Preview Text Direction", preview_text_direction)
    add_row(table, "Primary Message", primary_message)
    add_bullets(table, "Content Block Structure", content_blocks)
    add_row(table, "CTA Direction", cta_direction)
    add_row(table, "Tone Notes", tone_notes)

    doc.add_paragraph("")


# ─────────────────────────────────────────────
# BUILD DOC
# ─────────────────────────────────────────────
doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)

doc.add_heading("Rho Nutrition", 0)
p = doc.add_paragraph("Post Purchase Flow — Email Briefs v2")
p.runs[0].font.size = Pt(13)
p = doc.add_paragraph("DigitsUp x Rho Nutrition  |  All 6 Personas  |  Day 3: Welcome + How to Use")
p.runs[0].font.size = Pt(10)
p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
doc.add_paragraph("")


# ─── P1 EXECUTIVE ───
add_persona_brief(
    doc, "P1", "The High-Performance Executive",
    email_num=1, send_day=3,
    objective="Reinforce the purchase decision and frame daily usage as a performance protocol. Remove friction from starting. Position Rho as a high-ROI addition to an already optimized life.",
    subject_lines=[
        "Your optimization protocol starts now",
        "One serving. Every morning. Here is why it matters.",
        "You made a smart call. Here is how to get the most from it.",
    ],
    preview_text_direction="Short, efficient. Mirror their mindset. Example: 'Takes 30 seconds. Works all day.'",
    primary_message="You invested in your performance. Now activate it. One serving of Rho in the morning delivers NAD+ and liposomal nutrients at the cellular level. Faster than capsules. No wasted product. Frame as a protocol, not a supplement.",
    content_blocks=[
        "Hero: Confident welcome. 'Your optimization protocol starts now.' Acknowledge the purchase as a smart decision.",
        "Secondary: 3-step usage protocol. Formatted as steps. Clean, scannable. No extra language.",
        "Secondary body: One paragraph on liposomal delivery vs capsules. Frame as ROI — you absorb what you paid for.",
        "Tertiary: Short timeline. Week 1, 2, 4. Frame as compounding ROI. Ends with a reminder that consistency is the protocol.",
    ],
    cta_direction="Direct and action-oriented. 'Add to Your Morning Stack.' No urgency. No discount. Just the next logical step.",
    tone_notes="Boardroom confidence. Zero fluff. Short sentences. Every line earns its place. Think internal memo, not wellness newsletter."
)

# ─── P2 CREATIVE ───
add_persona_brief(
    doc, "P2", "The Creative Entrepreneur",
    email_num=1, send_day=3,
    objective="Welcome the customer into the Rho ritual. Frame daily usage as an intentional, elevated moment. Make starting feel inspiring, not clinical.",
    subject_lines=[
        "Your new ritual is ready",
        "The morning routine just got an upgrade",
        "Clean. Simple. Yours.",
    ],
    preview_text_direction="Warm and sensory. Example: 'A 30-second ritual that sets the tone for everything after.'",
    primary_message="You chose a product that fits the way you live. Clean ingredients, thoughtful formulation, nothing unnecessary. The liquid format is intentional. Position this as a small act of self-care that supports the work they do.",
    content_blocks=[
        "Hero: Aesthetic welcome. Clean, morning energy. Warm tone. Not clinical.",
        "Secondary: How to build the ritual. Sensory and simple. Focus on ease and intention.",
        "Secondary: Clean label callout. What is in it and what is not. Transparency as a point of pride.",
        "Tertiary: One short paragraph connecting mental clarity to clean nutrition. NAD+ and creative output.",
    ],
    cta_direction="Soft and inviting. 'Make It Your Morning Ritual.' Feels like an invitation, not an instruction.",
    tone_notes="Warm, considered, artisan. Reads like a brand they already love. No clinical language. No mass-market wellness. White space is your friend."
)

# ─── P3 EDUCATOR ───
add_persona_brief(
    doc, "P3", "The Community Educator and Parent",
    email_num=1, send_day=3,
    objective="Make the customer feel good about their decision. Reassuring usage instructions. Reinforce that this is a safe, trusted product built for real people with full lives.",
    subject_lines=[
        "You made a good choice for yourself",
        "Here is how to get started with Rho",
        "Simple steps for your new daily routine",
    ],
    preview_text_direction="Warm and affirming. Example: 'You deserve this. Here is how to make it work for your day.'",
    primary_message="Between students, family, and everything else they carry, this is one thing that is just for them. Lead with warmth. Reinforce safety, simplicity, and the founder mission. One serving. That is all.",
    content_blocks=[
        "Hero: Warm welcome. Acknowledge how much they give. This is for them.",
        "Secondary: Simple usage instructions. No jargon. 'One serving in the morning with or without food.'",
        "Secondary: Safety and trust block. Clean ingredients, no harmful additives, founder mission. 2-3 sentences.",
        "Tertiary: Gentle timeline. Steady energy and immune support. Framed around showing up for others.",
    ],
    cta_direction="Encouraging and low-pressure. 'Start Your Routine Today.' Feels like a nudge from a trusted friend.",
    tone_notes="Human, warm, reassuring. No science-heavy language. No performance framing. This persona responds to trust and community."
)

# ─── P4 CLINICAL ───
add_persona_brief(
    doc, "P4", "The Clinical Wellness Advocate",
    email_num=1, send_day=3,
    objective="Confirm the clinical rationale for their purchase. Provide precise dosing and timing guidance. Acknowledge their expertise. Confirm they made the right call.",
    subject_lines=[
        "Your liposomal protocol is active",
        "Dosing, timing, and what to expect from day one",
        "The formulation you chose. Here is how to use it correctly.",
    ],
    preview_text_direction="Clinical and direct. Example: 'Optimal absorption window, full formulation breakdown, and what to monitor.'",
    primary_message="They understand bioavailability. They chose liposomal for a reason. Confirm they were right and give them the clinical detail they expect: dosing, timing, absorption, key compounds. Peer-level throughout. No hand-holding.",
    content_blocks=[
        "Hero: Peer-level acknowledgment. 'You already know what liposomal means. Here is how to use Rho correctly.'",
        "Secondary: Dosing protocol table or formatted list. Timing, amount, fasting notes, absorption window.",
        "Secondary: Key active compounds. Glutathione, NAD+, Curcumin, Resveratrol. Delivery mechanism. One line each.",
        "Tertiary: Clinical timeline. Week 1, 2, 4. Framed at the cellular level, not symptomatic.",
    ],
    cta_direction="Evidence-driven. 'View the Full Formulation.' Invites them deeper without overselling.",
    tone_notes="Peer-level. Use correct terminology: bioavailability, liposomal, phospholipid, cellular absorption. Confident and precise. Never dumbed down. Never enthusiastic."
)

# ─── P5 PERFORMANCE ───
add_persona_brief(
    doc, "P5", "The Active Performance Optimizer",
    email_num=1, send_day=3,
    objective="Connect Rho directly to their training and recovery routine. Frame day one as the start of a performance system. Give them a clear, actionable stacking protocol.",
    subject_lines=[
        "Recovery starts with day one",
        "Here is how to stack Rho into your routine",
        "Day one. The protocol begins.",
    ],
    preview_text_direction="Action-oriented. Example: 'Morning timing, anti-inflammatory support, and what you should start noticing by week two.'",
    primary_message="Recovery is not what happens after training. It is the system built before the next session. Rho fits into that system. Morning timing, consistent daily use, Curcumin and Resveratrol doing their work. Give them the protocol.",
    content_blocks=[
        "Hero: Performance framing. 'Recovery does not start after the session. It starts here.'",
        "Secondary: Stacking protocol. When to take it, why morning timing matters for anti-inflammatory compounds.",
        "Secondary: Curcumin + Resveratrol callout. What they do, why bioavailability matters for performance.",
        "Tertiary: Week-by-week physical signals. What to notice and when. Validate what they will feel.",
    ],
    cta_direction="Active and peer-level. 'Add It to Your Stack.' Sounds like advice from a trusted coach, not a brand.",
    tone_notes="Knowledgeable, direct, energetic. Results and recovery are the only currencies. No lifestyle language. No ritual framing."
)

# ─── P6 TRADE ───
add_persona_brief(
    doc, "P6", "The Skilled Trade and Industry Leader",
    email_num=1, send_day=3,
    objective="Get straight to it. How to take it and what to expect. No story. No ritual. Practical instructions from someone who respects their time.",
    subject_lines=[
        "Here is how to get the most out of it",
        "Simple. Consistent. Effective. Your guide to Rho.",
        "You bought it. Here is exactly how to use it.",
    ],
    preview_text_direction="Blunt and practical. Example: 'One serving a day. Morning is best. Here is what consistent use does.'",
    primary_message="Take one serving in the morning. Do it every day. Explain liposomal vs standard in one paragraph using money language. Joint support and energy take weeks. No hype. Stick with it.",
    content_blocks=[
        "Hero: No-fluff welcome. 'You made a solid decision. Here is how to use it right.'",
        "Secondary: Usage instructions. Bullet format. When, how much, with or without food.",
        "Secondary: Why liposomal works. One short paragraph. Framed as: you are not wasting your money.",
        "Tertiary: Realistic timeline. Joint support and energy in weeks, not days. Honest and direct.",
    ],
    cta_direction="Blunt and direct. 'Start Today.' No decoration.",
    tone_notes="Straight talk. Short sentences. Zero lifestyle or wellness language. If it sounds like a wellness newsletter, rewrite it."
)

path = "/Users/minigrill/.openclaw/workspace/projects/outersignal/rho-nutrition/Rho Nutrition Post Purchase Briefs v2.docx"
doc.save(path)
print(f"Saved: {path}")
