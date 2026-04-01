from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'DDDDDD')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_block_header(doc, block_name, block_type):
    p = doc.add_paragraph()
    run = p.add_run(f"{block_name}  ")
    run.bold = True
    run.font.size = Pt(11)
    tag = p.add_run(f"[{block_type}]")
    tag.font.size = Pt(9)
    tag.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)

def add_field(doc, label, value, note=None):
    p = doc.add_paragraph()
    run = p.add_run(label + ": ")
    run.bold = True
    run.font.size = Pt(10)
    p.add_run(value).font.size = Pt(10)
    p.paragraph_format.space_after = Pt(3)
    if note:
        n = doc.add_paragraph()
        nr = n.add_run(f"Design note: {note}")
        nr.font.size = Pt(9)
        nr.font.italic = True
        nr.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        n.paragraph_format.space_after = Pt(6)

def add_persona_section(doc, persona_code, persona_name, subject_a, subject_b, preheader, hero_headline, hero_body, hero_cta, secondary_headline, secondary_body, secondary_cta, tertiary_headline, tertiary_body, tertiary_cta=None, tertiary_note=None):
    doc.add_heading(f"{persona_code}: {persona_name}", level=2)

    # Meta
    add_field(doc, "Subject Line A", subject_a)
    add_field(doc, "Subject Line B", subject_b)
    add_field(doc, "Preheader", preheader)
    doc.add_paragraph("")

    # Hero
    add_block_header(doc, "HERO BLOCK", "Image + Text + CTA")
    add_field(doc, "Headline", hero_headline, "Set in image by designer. Bold, large, brand font.")
    add_field(doc, "Supporting copy", hero_body, "Live text below hero image. 2-3 short sentences max.")
    add_field(doc, "CTA", hero_cta, "Button. Centered.")

    # Secondary
    add_block_header(doc, "SECONDARY BLOCK", "Content + CTA")
    add_field(doc, "Section headline", secondary_headline, "Can be live text or set in image header. Designer's call.")
    add_field(doc, "Body copy", secondary_body, "Live text. Short paragraphs or 3-step list. Keep tight.")
    add_field(doc, "CTA", secondary_cta, "Text link or button below content.")

    # Tertiary
    add_block_header(doc, "TERTIARY BLOCK", "Supporting Content")
    add_field(doc, "Section headline", tertiary_headline, tertiary_note or "Live text headline. Smaller weight than secondary.")
    add_field(doc, "Body copy", tertiary_body, "1-3 short sentences. No wall of text.")
    if tertiary_cta:
        add_field(doc, "CTA", tertiary_cta, "Optional. Text link preferred at this level.")

    add_divider(doc)


# ─────────────────────────────────────────────
# TITLE
# ─────────────────────────────────────────────
doc.add_heading("Rho Nutrition", 0)
p = doc.add_paragraph("Post Purchase Flow — Email 1 (Day 3) Copy Doc v1")
p.runs[0].font.size = Pt(14)
p = doc.add_paragraph("DigitsUp x Rho Nutrition  |  All 6 Personas  |  Welcome + How to Use")
p.runs[0].font.size = Pt(10)
p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
doc.add_paragraph("")

p = doc.add_paragraph()
r = p.add_run("Copy rules: ")
r.bold = True
p.add_run("No em dashes. No emoji in body copy. No urgency or discount language. No generic wellness buzzwords. Short sentences. White space. Each block must earn its place.")
p.runs[-1].font.size = Pt(10)
p.runs[-1].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
doc.add_paragraph("")


# ─────────────────────────────────────────────
# P1: EXECUTIVE
# ─────────────────────────────────────────────
add_persona_section(
    doc,
    persona_code="P1",
    persona_name="The High-Performance Executive",
    subject_a="Your optimization protocol starts now",
    subject_b="One serving. Every morning. Here is why it matters.",
    preheader="Takes 30 seconds. Works all day.",
    hero_headline="Your Optimization Protocol Starts Now",
    hero_body="You made a smart investment. Now activate it. One serving of Rho each morning delivers NAD+ and liposomal nutrients directly at the cellular level. Faster than capsules. No wasted product. No guesswork.",
    hero_cta="Add to Your Morning Stack",
    secondary_headline="How to Use It",
    secondary_body="Step 1: Take one serving first thing in the morning.\nStep 2: With or without food. Liquid format absorbs either way.\nStep 3: Repeat daily. Cellular benefits compound over time.\n\nThe liposomal delivery system means your body absorbs what you paid for. Standard capsule supplements lose a significant portion to digestion. Rho does not.",
    secondary_cta="View Full Formulation",
    tertiary_headline="What to Expect",
    tertiary_body="Week 1: Absorption begins. No dramatic change yet. That is normal.\nWeek 2: Cognitive clarity and sustained energy start to stabilize.\nWeek 4: Cellular repair compounds. This is where the ROI shows.\n\nConsistency is the protocol. Missing days resets the compounding.",
    tertiary_cta=None,
    tertiary_note="Live text. Formatted as a short timeline. Clean, scannable."
)

# ─────────────────────────────────────────────
# P2: CREATIVE
# ─────────────────────────────────────────────
add_persona_section(
    doc,
    persona_code="P2",
    persona_name="The Creative Entrepreneur",
    subject_a="Your new ritual is ready",
    subject_b="The morning routine just got an upgrade",
    preheader="A 30-second ritual that sets the tone for everything after.",
    hero_headline="Your New Ritual Is Ready",
    hero_body="You chose a product that fits the way you live. Clean ingredients. Thoughtful formulation. Nothing unnecessary. The liquid format is intentional: easier to absorb, easier to build into a morning you already care about.",
    hero_cta="Make It Your Morning Ritual",
    secondary_headline="How to Build the Ritual",
    secondary_body="Take one serving each morning. The liquid format is designed to fit into the moment between coffee and everything else. No mixing. No waiting.\n\nWhat is in it matters as much as what is not. No fillers. No synthetic binders. Just the compounds your body can actually use.\n\nThis is not about adding another thing to your routine. It is about making the routine better.",
    secondary_cta="See What Is Inside",
    tertiary_headline="Why Clean Formulation Matters for Creative Work",
    tertiary_body="Mental clarity is not separate from physical health. NAD+ supports cellular energy. Liposomal absorption means the ingredients reach your cells, not just your bloodstream. The work you do requires a mind that is clear and a body that can keep up. Rho is built for that.",
    tertiary_cta=None,
    tertiary_note="Live text. Soft, considered tone. Not clinical. Not hype."
)

# ─────────────────────────────────────────────
# P3: EDUCATOR/PARENT
# ─────────────────────────────────────────────
add_persona_section(
    doc,
    persona_code="P3",
    persona_name="The Community Educator and Parent",
    subject_a="You made a good choice for yourself",
    subject_b="Here is how to get started with Rho",
    preheader="You deserve this. Here is how to make it work for your day.",
    hero_headline="You Made a Good Choice for Yourself",
    hero_body="Between your students, your family, and everything else you carry, this is one thing that is just for you. Rho is clean, simple, and safe. One serving in the morning. That is all it takes to start.",
    hero_cta="Start Your Routine Today",
    secondary_headline="Getting Started Is Simple",
    secondary_body="Take one serving each morning, with or without food. The liquid format makes it easy to fit into a busy morning without any prep.\n\nRho is formulated without harmful additives, synthetic fillers, or anything you would not want in your home. The founders built this for their own families first. That standard has not changed.\n\nYou do not need to change anything about your day. Just add this to it.",
    secondary_cta="Learn What Is Inside",
    tertiary_headline="What Consistent Use Supports",
    tertiary_body="Immune function. Sustained energy. Cellular repair over time. These are not dramatic claims. They are what consistent, daily use of clean, bioavailable nutrients does. You show up for everyone else every day. This helps you keep doing that.",
    tertiary_cta=None,
    tertiary_note="Live text. Warm, grounding. No hyperbole."
)

# ─────────────────────────────────────────────
# P4: CLINICAL
# ─────────────────────────────────────────────
add_persona_section(
    doc,
    persona_code="P4",
    persona_name="The Clinical Wellness Advocate",
    subject_a="Your liposomal protocol is active",
    subject_b="Dosing, timing, and what to expect from day one",
    preheader="Optimal absorption window, full formulation breakdown, and what to monitor.",
    hero_headline="Your Liposomal Protocol Is Active",
    hero_body="You understand why phospholipid encapsulation outperforms standard delivery. Rho uses liposomal technology to protect active compounds through the digestive environment and deliver them at the cellular level. You made the right clinical call.",
    hero_cta="View the Full Formulation",
    secondary_headline="Dosing and Timing Protocol",
    secondary_body="Dose: One serving daily.\nTiming: Morning, ideally fasting or with a light meal for optimal absorption window.\nFormat: Liquid. No disintegration delay. Bioavailability significantly higher than encapsulated equivalents.\n\nKey compounds: Glutathione (reduced, liposomal), NAD+ precursors, Curcumin (liposomal), Resveratrol, Vitamin C.\n\nConsistency drives outcomes. Cellular repair is cumulative. Clinical benefit is not symptomatic in week one.",
    secondary_cta="Read the Research",
    tertiary_headline="Clinical Timeline",
    tertiary_body="Days 1 to 7: Absorption and baseline establishment. No significant symptomatic change expected.\nDays 8 to 14: Intracellular Glutathione levels begin to normalize. Cognitive and energy stabilization reported.\nDays 21 to 30: NAD+ pathway support accumulates. Inflammatory markers may begin to shift in patients with baseline elevation.\n\nFor patient recommendation purposes: 60-day minimum protocol recommended for measurable outcome assessment.",
    tertiary_cta=None,
    tertiary_note="Live text. Clinical format. Bullet or numbered list preferred. Peer-level language throughout."
)

# ─────────────────────────────────────────────
# P5: PERFORMANCE
# ─────────────────────────────────────────────
add_persona_section(
    doc,
    persona_code="P5",
    persona_name="The Active Performance Optimizer",
    subject_a="Recovery starts with day one",
    subject_b="Here is how to stack Rho into your routine",
    preheader="Morning timing, anti-inflammatory support, and what you should feel by week two.",
    hero_headline="Recovery Starts With Day One",
    hero_body="Recovery is not what happens after you train. It is the system you build before the next session. Rho fits into that system. Take it in the morning. Let the compounds work consistently. By day eight, you will notice the difference.",
    hero_cta="Add It to Your Stack",
    secondary_headline="How to Stack It",
    secondary_body="Take one serving every morning. Curcumin and Resveratrol are not acute compounds. They work through daily accumulation, not single doses. Morning timing allows the anti-inflammatory pathway to be active throughout your training day.\n\nLiposomal delivery means the Curcumin your body actually absorbs is significantly higher than standard supplement equivalents. You are not wasting product. You are loading the system.\n\nDo not skip days. Consistency is what separates results from wasted spend.",
    secondary_cta="See the Science Behind Recovery",
    tertiary_headline="What You Should Start Noticing",
    tertiary_body="By day 3 to 5: Nothing dramatic. That is normal. The compounds are building.\nBy day 8: Joint comfort during and after training should begin to shift. Less stiffness. Faster return to baseline.\nBy day 21: Recovery speed improves. Inflammation response to heavy training sessions decreases.\n\nIf you are tracking performance metrics, log your recovery quality daily. The pattern will show up.",
    tertiary_cta=None,
    tertiary_note="Live text. Coach-level voice. Practical and direct. No wellness softness."
)

# ─────────────────────────────────────────────
# P6: TRADE
# ─────────────────────────────────────────────
add_persona_section(
    doc,
    persona_code="P6",
    persona_name="The Skilled Trade and Industry Leader",
    subject_a="Here is how to get the most out of it",
    subject_b="Simple. Consistent. Effective.",
    preheader="One serving a day. Morning is best. Here is what consistent use does.",
    hero_headline="Here Is How to Get the Most Out of It",
    hero_body="You bought a quality product. Now use it right. One serving in the morning, every day. That is it. The compounds build up in your system over time. Skipping days slows the process. Staying consistent is what makes it work.",
    hero_cta="Start Today",
    secondary_headline="How to Take It",
    secondary_body="One serving. Every morning. With or without food.\n\nThe liquid format absorbs faster than capsules. That matters because most supplements lose a large portion of their active ingredients during digestion. Rho does not. Liposomal delivery gets the compounds where they need to go. That is why it costs more than what you find at a drugstore. And why it works when those do not.\n\nKeep it somewhere you will see it in the morning. Counter, fridge, wherever. Make it automatic.",
    secondary_cta="See Why It Is Different",
    tertiary_headline="What to Expect",
    tertiary_body="Week one: Not much. That is normal. The compounds are loading.\nWeek two: Joint comfort starts to improve. Especially if you are on your feet all day or doing physical work.\nWeek three and beyond: Sustained energy and reduced soreness. The physical wear your job puts on your body needs consistent support. This is what that looks like.\n\nStick with it.",
    tertiary_cta=None,
    tertiary_note="Live text. Blunt. No flourishes. Reads like a straight conversation."
)

# Save
path = "/Users/minigrill/.openclaw/workspace/projects/outersignal/rho-nutrition/Rho Nutrition Email 1 Day 3 Copy Doc v1.docx"
doc.save(path)
print(f"Saved: {path}")
