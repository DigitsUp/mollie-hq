from docx import Document
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_border(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '999999')
        tcBorders.append(border)
    tcPr.append(tcBorders)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_row(table, label, value, is_note=False):
    row = table.add_row()
    
    # Label cell
    lc = row.cells[0]
    set_cell_border(lc)
    set_cell_bg(lc, 'F5F5F5')
    p = lc.paragraphs[0]
    run = p.add_run(label)
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    lc.width = Inches(1.8)

    # Value cell
    vc = row.cells[1]
    set_cell_border(vc)
    
    if is_note:
        p = vc.paragraphs[0]
        run = p.add_run(value)
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        run.font.name = 'Calibri'
    else:
        # Handle multiline values
        lines = value.split('\n')
        first = True
        for line in lines:
            if first:
                p = vc.paragraphs[0]
                first = False
            else:
                p = vc.add_paragraph()
            
            if line.startswith('**') and line.endswith('**'):
                run = p.add_run(line[2:-2])
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
            elif line.startswith('- '):
                p.style = 'List Bullet'
                run = p.add_run(line[2:])
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
            else:
                run = p.add_run(line)
                run.font.size = Pt(10)
                run.font.name = 'Calibri'
            p.paragraph_format.space_after = Pt(2)

    return row

def add_section_header(table, title):
    row = table.add_row()
    row.cells[0].merge(row.cells[1])
    cell = row.cells[0]
    set_cell_bg(cell, '1A1A1A')
    set_cell_border(cell)
    p = cell.paragraphs[0]
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

def add_block_header(table, title, block_type):
    row = table.add_row()
    row.cells[0].merge(row.cells[1])
    cell = row.cells[0]
    set_cell_bg(cell, 'E8F0E8')
    set_cell_border(cell)
    p = cell.paragraphs[0]
    run = p.add_run(f"{title}  ")
    run.bold = True
    run.font.size = Pt(10)
    run.font.name = 'Calibri'
    tag = p.add_run(f"[{block_type}]")
    tag.font.size = Pt(9)
    tag.font.italic = True
    tag.font.color.rgb = RGBColor(0x55, 0x88, 0x55)
    tag.font.name = 'Calibri'
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

def add_persona(doc, persona_code, persona_name,
                subject_a, subject_b, preheader,
                hero_headline, hero_copy, hero_cta,
                secondary_headline, secondary_copy, secondary_cta,
                tertiary_headline, tertiary_copy, tertiary_cta=None):

    # Persona heading
    h = doc.add_heading(f"{persona_code}: {persona_name}", level=2)
    h.paragraph_format.space_before = Pt(20)
    h.paragraph_format.space_after = Pt(8)

    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    table.columns[0].width = Inches(1.8)
    table.columns[1].width = Inches(4.8)

    # Meta
    add_section_header(table, "Email Meta")
    add_row(table, "Subject Line A", subject_a)
    add_row(table, "Subject Line B", subject_b)
    add_row(table, "Preheader", preheader)

    # Hero
    add_block_header(table, "Hero Block", "Image + Text + CTA")
    add_row(table, "Headline", hero_headline)
    add_row(table, "Supporting Copy", hero_copy)
    add_row(table, "CTA", hero_cta)
    add_row(table, "Design Note", "Full-width hero image. Headline set in image by designer. Supporting copy is live text below image. CTA is centered button.", is_note=True)

    # Secondary
    add_block_header(table, "Secondary Block", "Content + CTA")
    add_row(table, "Section Headline", secondary_headline)
    add_row(table, "Body Copy", secondary_copy)
    add_row(table, "CTA", secondary_cta)
    add_row(table, "Design Note", "Section headline can be live text or set in light image header. Body copy is live text. Keep tight. Short paragraphs or numbered steps.", is_note=True)

    # Tertiary
    add_block_header(table, "Tertiary Block", "Supporting Content")
    add_row(table, "Section Headline", tertiary_headline)
    add_row(table, "Body Copy", tertiary_copy)
    if tertiary_cta:
        add_row(table, "CTA", tertiary_cta)
    add_row(table, "Design Note", "Live text throughout. Lighter visual weight than secondary. No competing with the primary CTA.", is_note=True)

    doc.add_paragraph("")


# ─────────────────────────────────────────────
# BUILD DOC
# ─────────────────────────────────────────────
doc = Document()

# Page margins
from docx.oxml import OxmlElement
section = doc.sections[0]
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.8)
section.left_margin = Inches(0.9)
section.right_margin = Inches(0.9)

# Title
doc.add_heading("Rho Nutrition", 0)
p = doc.add_paragraph("Post Purchase Flow — Email 1 (Day 3) Copy Doc v2")
p.runs[0].font.size = Pt(13)
p = doc.add_paragraph("DigitsUp x Rho Nutrition  |  All 6 Personas  |  Welcome + How to Use")
p.runs[0].font.size = Pt(10)
p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)
doc.add_paragraph("")

p = doc.add_paragraph()
r = p.add_run("Copy rules: ")
r.bold = True
r.font.size = Pt(10)
p.add_run("No em dashes. No emoji in body copy. No urgency or discount language. Short sentences. White space. No generic wellness buzzwords.")
p.runs[-1].font.size = Pt(10)
p.runs[-1].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_paragraph("")


# ─── P1 EXECUTIVE ───
add_persona(
    doc, "P1", "The High-Performance Executive",
    subject_a="Your optimization protocol starts now",
    subject_b="One serving. Every morning. Here is why it matters.",
    preheader="Takes 30 seconds. Works all day.",
    hero_headline="Your Optimization Protocol Starts Now",
    hero_copy="You made a smart investment. Now activate it. One serving of Rho each morning delivers NAD+ and liposomal nutrients directly at the cellular level. Faster than capsules. No wasted product. No guesswork.",
    hero_cta="Add to Your Morning Stack",
    secondary_headline="How to Use It",
    secondary_copy="Step 1: Take one serving first thing in the morning.\nStep 2: With or without food. Liquid format absorbs either way.\nStep 3: Repeat daily. Cellular benefits compound over time.\n\nThe liposomal delivery system means your body absorbs what you paid for. Standard capsule supplements lose a significant portion to digestion. Rho does not.",
    secondary_cta="View Full Formulation",
    tertiary_headline="What to Expect",
    tertiary_copy="Week 1: Absorption begins. No dramatic change yet. That is normal.\nWeek 2: Cognitive clarity and sustained energy start to stabilize.\nWeek 4: Cellular repair compounds. This is where the ROI shows.\n\nConsistency is the protocol. Missing days resets the compounding.",
)

# ─── P2 CREATIVE ───
add_persona(
    doc, "P2", "The Creative Entrepreneur",
    subject_a="Your new ritual is ready",
    subject_b="The morning routine just got an upgrade",
    preheader="A 30-second ritual that sets the tone for everything after.",
    hero_headline="Your New Ritual Is Ready",
    hero_copy="You chose a product that fits the way you live. Clean ingredients. Thoughtful formulation. Nothing unnecessary. The liquid format is intentional: easier to absorb, easier to build into a morning you already care about.",
    hero_cta="Make It Your Morning Ritual",
    secondary_headline="How to Build the Ritual",
    secondary_copy="Take one serving each morning. The liquid format is designed to fit into the moment between coffee and everything else. No mixing. No waiting.\n\nWhat is in it matters as much as what is not. No fillers. No synthetic binders. Just the compounds your body can actually use.\n\nThis is not about adding another thing to your routine. It is about making the routine better.",
    secondary_cta="See What Is Inside",
    tertiary_headline="Why Clean Formulation Matters for Creative Work",
    tertiary_copy="Mental clarity is not separate from physical health. NAD+ supports cellular energy. Liposomal absorption means the ingredients reach your cells, not just your bloodstream. The work you do requires a mind that is clear and a body that can keep up.",
)

# ─── P3 EDUCATOR ───
add_persona(
    doc, "P3", "The Community Educator and Parent",
    subject_a="You made a good choice for yourself",
    subject_b="Here is how to get started with Rho",
    preheader="You deserve this. Here is how to make it work for your day.",
    hero_headline="You Made a Good Choice for Yourself",
    hero_copy="Between your students, your family, and everything else you carry, this is one thing that is just for you. Rho is clean, simple, and safe. One serving in the morning. That is all it takes to start.",
    hero_cta="Start Your Routine Today",
    secondary_headline="Getting Started Is Simple",
    secondary_copy="Take one serving each morning, with or without food. The liquid format makes it easy to fit into a busy morning without any prep.\n\nRho is formulated without harmful additives, synthetic fillers, or anything you would not want in your home. The founders built this for their own families first. That standard has not changed.\n\nYou do not need to change anything about your day. Just add this to it.",
    secondary_cta="Learn What Is Inside",
    tertiary_headline="What Consistent Use Supports",
    tertiary_copy="Immune function. Sustained energy. Cellular repair over time. These are not dramatic claims. They are what consistent, daily use of clean, bioavailable nutrients does. You show up for everyone else every day. This helps you keep doing that.",
)

# ─── P4 CLINICAL ───
add_persona(
    doc, "P4", "The Clinical Wellness Advocate",
    subject_a="Your liposomal protocol is active",
    subject_b="Dosing, timing, and what to expect from day one",
    preheader="Optimal absorption window, full formulation breakdown, and what to monitor.",
    hero_headline="Your Liposomal Protocol Is Active",
    hero_copy="You understand why phospholipid encapsulation outperforms standard delivery. Rho uses liposomal technology to protect active compounds through the digestive environment and deliver them at the cellular level. You made the right clinical call.",
    hero_cta="View the Full Formulation",
    secondary_headline="Dosing and Timing Protocol",
    secondary_copy="Dose: One serving daily.\nTiming: Morning, ideally fasting or with a light meal for optimal absorption.\nFormat: Liquid. No disintegration delay. Bioavailability significantly higher than encapsulated equivalents.\n\nKey compounds: Glutathione (reduced, liposomal), NAD+ precursors, Curcumin (liposomal), Resveratrol, Vitamin C.\n\nConsistency drives outcomes. Cellular repair is cumulative. Clinical benefit is not symptomatic in week one.",
    secondary_cta="Read the Research",
    tertiary_headline="Clinical Timeline",
    tertiary_copy="Days 1 to 7: Absorption and baseline establishment. No significant symptomatic change expected.\nDays 8 to 14: Intracellular Glutathione levels begin to normalize. Cognitive and energy stabilization reported.\nDays 21 to 30: NAD+ pathway support accumulates. Inflammatory markers may begin to shift.\n\nFor patient recommendation: 60-day minimum protocol recommended for measurable outcome assessment.",
)

# ─── P5 PERFORMANCE ───
add_persona(
    doc, "P5", "The Active Performance Optimizer",
    subject_a="Recovery starts with day one",
    subject_b="Here is how to stack Rho into your routine",
    preheader="Morning timing, anti-inflammatory support, and what you should feel by week two.",
    hero_headline="Recovery Starts With Day One",
    hero_copy="Recovery is not what happens after you train. It is the system you build before the next session. Rho fits into that system. Take it in the morning. Let the compounds work consistently. By day eight, you will notice the difference.",
    hero_cta="Add It to Your Stack",
    secondary_headline="How to Stack It",
    secondary_copy="Take one serving every morning. Curcumin and Resveratrol are not acute compounds. They work through daily accumulation, not single doses. Morning timing allows the anti-inflammatory pathway to be active throughout your training day.\n\nLiposomal delivery means the Curcumin your body actually absorbs is significantly higher than standard supplement equivalents. You are not wasting product. You are loading the system.\n\nDo not skip days. Consistency is what separates results from wasted spend.",
    secondary_cta="See the Science Behind Recovery",
    tertiary_headline="What You Should Start Noticing",
    tertiary_copy="Days 3 to 5: Nothing dramatic. That is normal. The compounds are building.\nDay 8: Joint comfort during and after training should begin to shift. Less stiffness. Faster return to baseline.\nDay 21: Recovery speed improves. Inflammation response to heavy training sessions decreases.\n\nIf you are tracking performance metrics, log your recovery quality daily. The pattern will show up.",
)

# ─── P6 TRADE ───
add_persona(
    doc, "P6", "The Skilled Trade and Industry Leader",
    subject_a="Here is how to get the most out of it",
    subject_b="Simple. Consistent. Effective.",
    preheader="One serving a day. Morning is best. Here is what consistent use does.",
    hero_headline="Here Is How to Get the Most Out of It",
    hero_copy="You bought a quality product. Now use it right. One serving in the morning, every day. That is it. The compounds build up in your system over time. Skipping days slows the process. Staying consistent is what makes it work.",
    hero_cta="Start Today",
    secondary_headline="How to Take It",
    secondary_copy="One serving. Every morning. With or without food.\n\nThe liquid format absorbs faster than capsules. Most supplements lose a large portion of their active ingredients during digestion. Rho does not. Liposomal delivery gets the compounds where they need to go. That is why it works when others do not.\n\nKeep it somewhere you will see it in the morning. Counter, fridge, wherever. Make it automatic.",
    secondary_cta="See Why It Is Different",
    tertiary_headline="What to Expect",
    tertiary_copy="Week one: Not much. That is normal. The compounds are loading.\nWeek two: Joint comfort starts to improve. Especially if you are on your feet all day or doing physical work.\nWeek three onward: Sustained energy and reduced soreness. The physical wear your job puts on your body needs consistent support. This is what that looks like.\n\nStick with it.",
)

path = "/Users/minigrill/.openclaw/workspace/projects/outersignal/rho-nutrition/Rho Nutrition Email 1 Day 3 Copy Doc v2.docx"
doc.save(path)
print(f"Saved: {path}")
